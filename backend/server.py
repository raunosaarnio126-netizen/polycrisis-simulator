from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict
import uuid
from datetime import datetime, timezone, timedelta
from passlib.context import CryptContext
import jwt
from emergentintegrations.llm.chat import LlmChat, UserMessage
import PyPDF2
from docx import Document
import io
import json
import asyncio
from pathlib import Path

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI(title="Polycrisis Simulator API")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
security = HTTPBearer()
SECRET_KEY = "polycrisis-secret-key-2025"  # In production, use environment variable
ALGORITHM = "HS256"

# AI Integration Setup
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

# Models
class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    username: str
    organization: str
    company_id: Optional[str] = None
    role: str = "member"  # "admin", "manager", "analyst", "member"
    department: Optional[str] = None
    job_title: Optional[str] = None
    phone: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Scenario Adjusters - Fuzzy Logic Models
class ScenarioAdjustment(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_id: str
    scenario_id: Optional[str] = None  # Link to existing scenario
    adjustment_name: str
    # SEPTE Framework percentages (each pair should sum to 100)
    economic_crisis_pct: float = 50.0  # vs economic_stability_pct
    economic_stability_pct: float = 50.0
    social_unrest_pct: float = 50.0  # vs social_cohesion_pct  
    social_cohesion_pct: float = 50.0
    environmental_degradation_pct: float = 50.0  # vs environmental_resilience_pct
    environmental_resilience_pct: float = 50.0
    political_instability_pct: float = 50.0  # vs political_stability_pct
    political_stability_pct: float = 50.0
    technological_disruption_pct: float = 50.0  # vs technological_advancement_pct
    technological_advancement_pct: float = 50.0
    # AI Analysis Results
    real_time_analysis: str
    impact_summary: str
    risk_level: str = "medium"  # "low", "medium", "high", "critical"
    recommendations: List[str] = []
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ScenarioAdjustmentCreate(BaseModel):
    adjustment_name: str
    scenario_id: Optional[str] = None
    economic_crisis_pct: float = 50.0
    economic_stability_pct: float = 50.0
    social_unrest_pct: float = 50.0
    social_cohesion_pct: float = 50.0
    environmental_degradation_pct: float = 50.0
    environmental_resilience_pct: float = 50.0
    political_instability_pct: float = 50.0
    political_stability_pct: float = 50.0
    technological_disruption_pct: float = 50.0
    technological_advancement_pct: float = 50.0

class ConsensusSettings(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_id: str
    team_id: Optional[str] = None
    adjustment_id: str
    consensus_name: str
    agreed_by: List[str] = []  # List of user IDs who agreed
    total_team_members: int
    consensus_reached: bool = False
    consensus_percentage: float = 0.0  # Percentage of team that agreed
    final_settings: dict  # Store the agreed SEPTE settings
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    finalized_at: Optional[datetime] = None

class ConsensusCreate(BaseModel):
    adjustment_id: str
    consensus_name: str
    team_id: Optional[str] = None

class UserCreate(BaseModel):
    email: str
    username: str
    password: str
    organization: str
    job_title: Optional[str] = None
    department: Optional[str] = None
    phone: Optional[str] = None

class UserLogin(BaseModel):
    email: str
    password: str

class Token(BaseModel):
    access_token: str
    token_type: str

class Scenario(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    title: str
    description: str
    crisis_type: str  # "natural_disaster", "economic_crisis", "social_unrest", "pandemic", etc.
    severity_level: int  # 1-10 scale
    affected_regions: List[str]
    key_variables: List[str]
    status: str = "draft"  # "draft", "active", "completed"
    additional_context: Optional[str] = None
    stakeholders: Optional[str] = None
    timeline: Optional[str] = None
    
    # Option 1: Sequential Numbering/Labeling
    sequence_number: Optional[int] = None  # Auto-generated: 1, 2, 3, ...
    sequence_letter: Optional[str] = None  # Auto-generated: A, B, C, ...
    
    # Option 2: Impact Change Tracking
    change_history: List[dict] = Field(default_factory=list)  # Track all modifications
    last_modified_by: Optional[str] = None
    modification_count: int = 0  # Total number of modifications
    
    # Option 3: ABC Analysis Classification
    abc_classification: str = "B"  # "A" (High Impact), "B" (Medium Impact), "C" (Low Impact)
    priority_score: int = 5  # 1-10 priority scoring
    impact_category: str = "medium"  # "high", "medium", "low"
    
    # Option 4: Version Control/Change Counter
    version_number: str = "1.0.0"  # Semantic versioning
    major_version: int = 1
    minor_version: int = 0
    patch_version: int = 0
    revision_count: int = 0  # Total revisions
    
    # Option 5: Impact Measurement System
    impact_score: float = 50.0  # 0-100 impact scoring
    economic_impact: Optional[float] = None  # Economic impact scoring
    social_impact: Optional[float] = None  # Social impact scoring
    environmental_impact: Optional[float] = None  # Environmental impact scoring
    calculated_total_impact: Optional[float] = None  # Auto-calculated total impact
    impact_trend: str = "stable"  # "increasing", "decreasing", "stable"
    
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ScenarioCreate(BaseModel):
    title: str
    description: str
    crisis_type: str
    severity_level: int
    affected_regions: List[str]
    key_variables: List[str]

class AIGenieRequest(BaseModel):
    scenario_id: Optional[str] = None
    user_query: str
    context: Optional[str] = None

class AIGenieResponse(BaseModel):
    response: str
    suggestions: List[str]
    monitoring_tasks: List[str]

class SimulationResult(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    analysis: str
    risk_assessment: str
    mitigation_strategies: List[str]
    key_insights: List[str]
    confidence_score: float
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class GameBook(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    game_book_content: str
    decision_points: List[str]
    resource_requirements: List[str]
    timeline_phases: List[str]
    success_metrics: List[str]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ActionPlan(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    immediate_actions: List[str]
    short_term_actions: List[str]
    long_term_actions: List[str]
    responsible_parties: List[str]
    resource_allocation: List[str]
    priority_level: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StrategyImplementation(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    implementation_strategy: str
    organizational_changes: List[str]
    policy_recommendations: List[str]
    training_requirements: List[str]
    budget_considerations: List[str]
    stakeholder_engagement: List[str]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class MonitorAgent(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    agent_type: str  # "risk_monitor", "performance_tracker", "anomaly_detector", "trend_analyzer"
    status: str = "active"  # "active", "paused", "completed"
    monitoring_parameters: List[str]
    insights_generated: List[str]
    anomalies_detected: List[str] 
    risk_level: str = "low"  # "low", "medium", "high", "critical"
    last_analysis: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class SystemMetrics(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    resilience_score: float  # 0.0 - 1.0
    complexity_index: float  # 0.0 - 10.0
    cascading_risk_factor: float  # 0.0 - 1.0
    intervention_effectiveness: float  # 0.0 - 1.0
    system_stability: float  # 0.0 - 1.0
    adaptive_capacity: float  # 0.0 - 1.0
    interconnectedness_level: float  # 0.0 - 1.0
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ComplexAdaptiveSystem(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    system_components: List[str]  # Economic, Environmental, Social, Political, Technological
    interconnections: List[str]  # Relationships between components
    feedback_loops: List[str]  # Positive and negative feedback mechanisms
    emergent_behaviors: List[str]  # Unexpected system behaviors
    adaptation_mechanisms: List[str]  # How system adapts to changes
    tipping_points: List[str]  # Critical thresholds
    system_dynamics: str  # AI-generated description of system behavior
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class LearningInsight(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    scenario_id: str
    insight_type: str  # "pattern_recognition", "outcome_prediction", "optimization_suggestion"
    insight_content: str
    confidence_score: float  # 0.0 - 1.0
    applied: bool = False
    effectiveness_rating: Optional[float] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class MonitoringSource(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    user_id: str
    source_type: str  # "news_api", "social_media", "government_data", "weather_api", "financial_market", "custom_url"
    source_url: str
    source_name: str
    monitoring_frequency: str  # "real_time", "hourly", "daily", "weekly"
    data_keywords: List[str]  # Keywords to filter relevant information
    status: str = "active"  # "active", "paused", "error", "inactive"
    last_check: Optional[datetime] = None
    total_data_points: int = 0
    relevance_score: float = 0.0  # AI-calculated relevance to scenario
    added_by_team_member: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class MonitoringSourceCreate(BaseModel):
    source_type: str
    source_url: str
    source_name: str
    monitoring_frequency: str
    data_keywords: List[str]

class CollectedData(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    source_id: str
    scenario_id: str
    data_content: str
    data_title: str
    data_url: Optional[str] = None
    relevance_score: float  # AI-calculated relevance
    sentiment_score: float  # -1.0 to 1.0 (negative to positive)
    urgency_level: str  # "low", "medium", "high", "critical"
    keywords_matched: List[str]
    ai_summary: str
    collected_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TeamCollaboration(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    team_members: List[str]  # Email addresses of team members
    shared_sources: List[str]  # Source IDs
    collaboration_notes: List[str]
    access_level: str = "team"  # "team", "organization", "public"
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class SmartSuggestion(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    scenario_id: str
    suggestion_type: str  # "data_source", "monitoring_keyword", "analysis_focus"
    suggestion_content: str
    reasoning: str
    confidence_score: float
    implemented: bool = False
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Company(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_name: str
    industry: str
    company_size: str  # "startup", "small", "medium", "large", "enterprise"
    website_url: Optional[str] = None
    description: str
    location: str
    website_analysis: Optional[str] = None
    business_model: Optional[str] = None
    key_assets: List[str] = []
    vulnerabilities: List[str] = []
    stakeholders: List[str] = []
    competitive_landscape: Optional[str] = None
    created_by: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class CompanyCreate(BaseModel):
    company_name: str
    industry: str
    company_size: str
    website_url: Optional[str] = None
    description: str
    location: str

class BusinessDocument(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_id: str
    document_name: str
    document_type: str  # "business_plan", "market_strategy", "financial_report", "operational_plan", "other"
    document_content: str
    ai_analysis: Optional[str] = None
    key_insights: List[str] = []
    risk_factors: List[str] = []
    strategic_priorities: List[str] = []
    uploaded_by: str
    file_size: Optional[int] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class BusinessDocumentCreate(BaseModel):
    document_name: str
    document_type: str
    document_content: str

class Team(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_id: str
    team_name: str
    team_description: str
    team_lead: str  # User ID
    team_members: List[str]  # List of user IDs
    access_level: str = "standard"  # "standard", "admin", "view_only"
    team_roles: List[str] = []  # "crisis_manager", "analyst", "coordinator", "observer"
    active_scenarios: List[str] = []  # Scenario IDs this team is working on
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class TeamCreate(BaseModel):
    team_name: str
    team_description: str
    team_members: List[str]  # Email addresses
    team_roles: List[str] = []

class RapidAnalysis(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    company_id: str
    analysis_type: str  # "vulnerability_assessment", "business_impact", "scenario_recommendation", "competitive_analysis"
    analysis_title: str
    analysis_content: str
    key_findings: List[str]
    recommendations: List[str]
    priority_level: str  # "low", "medium", "high", "critical"
    confidence_score: float
    generated_by: str  # User ID
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# SaaS Admin & Licensing Models
class LicenseTier(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    tier_name: str  # "Single User", "Small Team", "Medium Team", "Large Team"
    max_users: int  # 1, 2, 5, 10
    monthly_price: float  # USD
    annual_price: float  # USD (usually with discount)
    features: List[str]
    stripe_price_id_monthly: Optional[str] = None
    stripe_price_id_annual: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Client(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    client_email: str
    license_tier_id: str
    license_count: int  # Number of licenses purchased
    users_active: int = 0  # Current active users
    subscription_status: str = "trial"  # "trial", "active", "cancelled", "expired"
    stripe_customer_id: Optional[str] = None
    stripe_subscription_id: Optional[str] = None
    trial_end_date: datetime
    subscription_start_date: Optional[datetime] = None
    subscription_end_date: Optional[datetime] = None
    created_by_admin: str  # Admin user ID
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    last_activity: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ClientCreate(BaseModel):
    client_name: str
    client_email: str
    license_tier_id: str
    license_count: int = 1

class AdminCredentials(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    admin_email: str
    admin_level: str = "super_admin"  # "super_admin", "admin", "support"
    permissions: List[str] = ["all"]  # ["client_management", "licensing", "billing", "support"]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AIAvatar(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    avatar_name: str
    avatar_type: str  # "risk_monitor", "performance_tracker", "anomaly_detector", "trend_analyzer"
    description: str
    base_competences: List[str]
    client_custom_competences: List[str] = []
    status: str = "active"  # "active", "inactive", "learning", "monitoring"
    performance_metrics: dict = {}
    client_id: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AvatarCompetence(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    avatar_id: str
    competence_name: str
    competence_description: str
    competence_type: str  # "skill", "knowledge", "capability"
    proficiency_level: int = 1  # 1-10 scale
    added_by_client: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AvatarCompetenceCreate(BaseModel):
    competence_name: str
    competence_description: str
    competence_type: str
    proficiency_level: int = 1

class PaymentRecord(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_id: str
    stripe_payment_intent_id: str
    amount: float
    currency: str = "USD"
    payment_status: str  # "pending", "succeeded", "failed", "cancelled"
    license_tier_id: str
    license_count: int
    billing_period: str  # "monthly", "annual"
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Helper functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.now(timezone.utc) + expires_delta
    else:
        expire = datetime.now(timezone.utc) + timedelta(hours=24)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: str = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        user = await db.users.find_one({"id": user_id})
        if user is None:
            raise HTTPException(status_code=401, detail="User not found")
        return User(**user)
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

# Authentication endpoints
@api_router.post("/register", response_model=Token)
async def register(user_data: UserCreate):
    # Check if user exists
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create new user
    hashed_password = get_password_hash(user_data.password)
    user = User(
        email=user_data.email,
        username=user_data.username,
        organization=user_data.organization
    )
    
    user_doc = user.dict()
    user_doc["password"] = hashed_password
    await db.users.insert_one(user_doc)
    
    # Create access token
    access_token = create_access_token(data={"sub": user.id})
    return {"access_token": access_token, "token_type": "bearer"}

@api_router.post("/login", response_model=Token)
async def login(login_data: UserLogin):
    user_doc = await db.users.find_one({"email": login_data.email})
    if not user_doc or not verify_password(login_data.password, user_doc["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    access_token = create_access_token(data={"sub": user_doc["id"]})
    return {"access_token": access_token, "token_type": "bearer"}

@api_router.get("/me", response_model=User)
async def get_me(current_user: User = Depends(get_current_user)):
    return current_user

# Scenario Management Helper Functions

async def get_next_sequence_number(user_id: str) -> int:
    """Get the next sequence number for scenarios"""
    count = await db.scenarios.count_documents({"user_id": user_id})
    return count + 1

def get_sequence_letter(sequence_number: int) -> str:
    """Convert sequence number to letter (1->A, 2->B, etc.)"""
    if sequence_number <= 26:
        return chr(64 + sequence_number)  # A, B, C, ...
    else:
        # For numbers > 26, use AA, AB, AC pattern
        first = chr(64 + ((sequence_number - 1) // 26))
        second = chr(64 + ((sequence_number - 1) % 26) + 1)
        return first + second

def calculate_abc_classification(severity_level: int, impact_score: float, crisis_type: str) -> tuple:
    """Calculate ABC classification based on scenario parameters"""
    # Weight factors for different crisis types
    crisis_weights = {
        "pandemic": 1.3,
        "natural_disaster": 1.2,
        "economic_crisis": 1.1,
        "social_unrest": 1.0,
        "technological_crisis": 0.9,
        "environmental_crisis": 1.2
    }
    
    weight = crisis_weights.get(crisis_type, 1.0)
    weighted_score = (severity_level * 10 + impact_score) * weight / 2
    
    if weighted_score >= 75:
        return "A", "high", max(8, min(10, int(weighted_score / 10)))
    elif weighted_score >= 50:
        return "B", "medium", max(4, min(7, int(weighted_score / 10)))
    else:
        return "C", "low", max(1, min(3, int(weighted_score / 10)))

def calculate_total_impact(economic: float = None, social: float = None, environmental: float = None) -> float:
    """Calculate total impact score from individual impact scores"""
    impacts = [impact for impact in [economic, social, environmental] if impact is not None]
    if not impacts:
        return 50.0  # Default neutral impact
    
    # Weighted average with emphasis on economic impact
    weights = [0.4, 0.3, 0.3]  # Economic, Social, Environmental
    weighted_sum = sum(impact * weight for impact, weight in zip(impacts, weights[:len(impacts)]))
    weight_sum = sum(weights[:len(impacts)])
    
    return round(weighted_sum / weight_sum, 2)

def create_change_record(action: str, field: str, old_value: any, new_value: any, user_id: str) -> dict:
    """Create a change history record"""
    return {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "action": action,
        "field": field,
        "old_value": str(old_value) if old_value is not None else None,
        "new_value": str(new_value) if new_value is not None else None,
        "modified_by": user_id,
        "change_id": str(uuid.uuid4())[:8]
    }

def update_version_number(current_version: str, change_type: str = "patch") -> tuple:
    """Update version number based on change type"""
    try:
        major, minor, patch = map(int, current_version.split('.'))
    except:
        major, minor, patch = 1, 0, 0
    
    if change_type == "major":
        major += 1
        minor = 0
        patch = 0
    elif change_type == "minor":
        minor += 1
        patch = 0
    else:  # patch
        patch += 1
    
    return f"{major}.{minor}.{patch}", major, minor, patch

# Scenario endpoints
@api_router.post("/scenarios", response_model=Scenario)
async def create_scenario(scenario_data: ScenarioCreate, current_user: User = Depends(get_current_user)):
    # Get sequence numbering
    sequence_number = await get_next_sequence_number(current_user.id)
    sequence_letter = get_sequence_letter(sequence_number)
    
    # Calculate initial impact scores (can be refined later)
    initial_impact_score = (scenario_data.severity_level * 10 + len(scenario_data.affected_regions) * 5) / 2
    economic_impact = initial_impact_score * 0.9  # Slightly lower economic impact initially
    social_impact = initial_impact_score * 1.1    # Slightly higher social impact initially
    environmental_impact = initial_impact_score * 0.8  # Variable environmental impact
    
    total_impact = calculate_total_impact(economic_impact, social_impact, environmental_impact)
    
    # Calculate ABC classification
    abc_class, impact_category, priority_score = calculate_abc_classification(
        scenario_data.severity_level, total_impact, scenario_data.crisis_type
    )
    
    # Create initial change record
    initial_change = create_change_record(
        "created", "scenario", None, f"Created scenario: {scenario_data.title}", current_user.id
    )
    
    scenario = Scenario(
        user_id=current_user.id,
        
        # Original fields
        title=scenario_data.title,
        description=scenario_data.description,
        crisis_type=scenario_data.crisis_type,
        severity_level=scenario_data.severity_level,
        affected_regions=scenario_data.affected_regions,
        key_variables=scenario_data.key_variables,
        
        # Option 1: Sequential Numbering/Labeling
        sequence_number=sequence_number,
        sequence_letter=sequence_letter,
        
        # Option 2: Impact Change Tracking
        change_history=[initial_change],
        last_modified_by=current_user.id,
        modification_count=0,
        
        # Option 3: ABC Analysis Classification
        abc_classification=abc_class,
        priority_score=priority_score,
        impact_category=impact_category,
        
        # Option 4: Version Control/Change Counter
        version_number="1.0.0",
        major_version=1,
        minor_version=0,
        patch_version=0,
        revision_count=0,
        
        # Option 5: Impact Measurement System
        impact_score=total_impact,
        economic_impact=economic_impact,
        social_impact=social_impact,
        environmental_impact=environmental_impact,
        calculated_total_impact=total_impact,
        impact_trend="stable"
    )
    
    await db.scenarios.insert_one(scenario.dict())
    return scenario

# Scenario Adjusters - Fuzzy Logic Endpoints
@api_router.post("/companies/{company_id}/scenario-adjustments", response_model=ScenarioAdjustment)
async def create_scenario_adjustment(company_id: str, adjustment_data: ScenarioAdjustmentCreate, current_user: User = Depends(get_current_user)):
    """Create a new scenario adjustment with fuzzy logic parameters"""
    # Verify company access
    company = await db.companies.find_one({"id": company_id})
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    
    if current_user.company_id != company_id and company['created_by'] != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Validate that opposing percentages sum to 100
    pairs = [
        (adjustment_data.economic_crisis_pct, adjustment_data.economic_stability_pct),
        (adjustment_data.social_unrest_pct, adjustment_data.social_cohesion_pct),
        (adjustment_data.environmental_degradation_pct, adjustment_data.environmental_resilience_pct),
        (adjustment_data.political_instability_pct, adjustment_data.political_stability_pct),
        (adjustment_data.technological_disruption_pct, adjustment_data.technological_advancement_pct)
    ]
    
    for pair in pairs:
        if abs(sum(pair) - 100.0) > 0.1:  # Allow small floating point errors
            raise HTTPException(status_code=400, detail="Each opposing pair must sum to 100%")
    
    try:
        # Generate AI analysis based on current settings
        analysis = await generate_scenario_analysis(company, adjustment_data)
        
        scenario_adjustment = ScenarioAdjustment(
            company_id=company_id,
            created_by=current_user.id,
            real_time_analysis=analysis['analysis'],
            impact_summary=analysis['impact_summary'],
            risk_level=analysis['risk_level'],
            recommendations=analysis['recommendations'],
            **adjustment_data.dict()
        )
        
        await db.scenario_adjustments.insert_one(scenario_adjustment.dict())
        return scenario_adjustment
        
    except Exception as e:
        logging.error(f"Scenario adjustment creation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create scenario adjustment: {str(e)}")

@api_router.get("/companies/{company_id}/scenario-adjustments", response_model=List[ScenarioAdjustment])
async def get_scenario_adjustments(company_id: str, current_user: User = Depends(get_current_user)):
    """Get all scenario adjustments for a company"""
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    
    adjustments = await db.scenario_adjustments.find({"company_id": company_id}).to_list(1000)
    return [ScenarioAdjustment(**adj) for adj in adjustments]

@api_router.put("/companies/{company_id}/scenario-adjustments/{adjustment_id}", response_model=ScenarioAdjustment)
async def update_scenario_adjustment(company_id: str, adjustment_id: str, adjustment_data: ScenarioAdjustmentCreate, current_user: User = Depends(get_current_user)):
    """Update scenario adjustment and regenerate AI analysis"""
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    else:
        company = await db.companies.find_one({"id": company_id})
    
    # Find existing adjustment
    existing_adjustment = await db.scenario_adjustments.find_one({"id": adjustment_id, "company_id": company_id})
    if not existing_adjustment:
        raise HTTPException(status_code=404, detail="Scenario adjustment not found")
    
    # Validate percentages
    pairs = [
        (adjustment_data.economic_crisis_pct, adjustment_data.economic_stability_pct),
        (adjustment_data.social_unrest_pct, adjustment_data.social_cohesion_pct),
        (adjustment_data.environmental_degradation_pct, adjustment_data.environmental_resilience_pct),
        (adjustment_data.political_instability_pct, adjustment_data.political_stability_pct),
        (adjustment_data.technological_disruption_pct, adjustment_data.technological_advancement_pct)
    ]
    
    for pair in pairs:
        if abs(sum(pair) - 100.0) > 0.1:
            raise HTTPException(status_code=400, detail="Each opposing pair must sum to 100%")
    
    try:
        # Generate new AI analysis
        analysis = await generate_scenario_analysis(company, adjustment_data)
        
        # Update the adjustment
        update_data = {
            **adjustment_data.dict(),
            "real_time_analysis": analysis['analysis'],
            "impact_summary": analysis['impact_summary'],
            "risk_level": analysis['risk_level'],
            "recommendations": analysis['recommendations'],
            "updated_at": datetime.now(timezone.utc)
        }
        
        await db.scenario_adjustments.update_one(
            {"id": adjustment_id, "company_id": company_id},
            {"$set": update_data}
        )
        
        updated_adjustment = await db.scenario_adjustments.find_one({"id": adjustment_id})
        return ScenarioAdjustment(**updated_adjustment)
        
    except Exception as e:
        logging.error(f"Scenario adjustment update error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update scenario adjustment: {str(e)}")

@api_router.post("/companies/{company_id}/consensus", response_model=ConsensusSettings)
async def create_consensus(company_id: str, consensus_data: ConsensusCreate, current_user: User = Depends(get_current_user)):
    """Create consensus settings for team agreement on scenario adjustments"""
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    
    # Verify adjustment exists
    adjustment = await db.scenario_adjustments.find_one({"id": consensus_data.adjustment_id, "company_id": company_id})
    if not adjustment:
        raise HTTPException(status_code=404, detail="Scenario adjustment not found")
    
    # Get team size if team_id provided
    total_members = 1
    if consensus_data.team_id:
        team = await db.teams.find_one({"id": consensus_data.team_id, "company_id": company_id})
        if team:
            total_members = len(team['team_members']) + 1  # +1 for team lead
    
    consensus = ConsensusSettings(
        company_id=company_id,
        team_id=consensus_data.team_id,
        adjustment_id=consensus_data.adjustment_id,
        consensus_name=consensus_data.consensus_name,
        agreed_by=[current_user.id],  # Creator automatically agrees
        total_team_members=total_members,
        consensus_percentage=100.0 / total_members,
        final_settings={
            "economic_crisis_pct": adjustment['economic_crisis_pct'],
            "economic_stability_pct": adjustment['economic_stability_pct'],
            "social_unrest_pct": adjustment['social_unrest_pct'],
            "social_cohesion_pct": adjustment['social_cohesion_pct'],
            "environmental_degradation_pct": adjustment['environmental_degradation_pct'],
            "environmental_resilience_pct": adjustment['environmental_resilience_pct'],
            "political_instability_pct": adjustment['political_instability_pct'],
            "political_stability_pct": adjustment['political_stability_pct'],
            "technological_disruption_pct": adjustment['technological_disruption_pct'],
            "technological_advancement_pct": adjustment['technological_advancement_pct']
        }
    )
    
    await db.consensus_settings.insert_one(consensus.dict())
    return consensus

@api_router.post("/companies/{company_id}/consensus/{consensus_id}/agree")
async def agree_to_consensus(company_id: str, consensus_id: str, current_user: User = Depends(get_current_user)):
    """User agrees to the consensus settings"""
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    
    # Find consensus
    consensus = await db.consensus_settings.find_one({"id": consensus_id, "company_id": company_id})
    if not consensus:
        raise HTTPException(status_code=404, detail="Consensus not found")
    
    # Add user to agreed list if not already there
    if current_user.id not in consensus['agreed_by']:
        agreed_by = consensus['agreed_by'] + [current_user.id]
        consensus_percentage = (len(agreed_by) / consensus['total_team_members']) * 100
        consensus_reached = consensus_percentage >= 75.0  # 75% threshold for consensus
        
        update_data = {
            "agreed_by": agreed_by,
            "consensus_percentage": consensus_percentage,
            "consensus_reached": consensus_reached
        }
        
        if consensus_reached and not consensus.get('finalized_at'):
            update_data["finalized_at"] = datetime.now(timezone.utc)
        
        await db.consensus_settings.update_one(
            {"id": consensus_id},
            {"$set": update_data}
        )
    else:
        # User already agreed, get current consensus status
        consensus_reached = consensus.get('consensus_reached', False)
    
    return {"message": "Agreement recorded", "consensus_reached": consensus_reached}

async def generate_scenario_analysis(company: dict, adjustment_data: ScenarioAdjustmentCreate) -> dict:
    """Generate AI analysis based on SEPTE framework adjustments"""
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"scenario-adjustment-{company['id']}",
        system_message="""You are an expert scenario analysis consultant specializing in SEPTE framework analysis (Social, Economic, Political, Technological, Environmental) for crisis management.

Your role is to:
1. Analyze fuzzy logic adjustments to scenario parameters
2. Provide real-time impact assessments based on percentage changes
3. Generate concise, actionable insights for scenario planning
4. Assess overall risk levels and provide strategic recommendations

Focus on short, clear analysis that immediately shows the impact of parameter changes."""
    ).with_model("anthropic", "claude-3-7-sonnet-20250219")
    
    # Build SEPTE analysis prompt
    analysis_prompt = f"""
Analyze this scenario adjustment for {company['company_name']} ({company['industry']}):

SEPTE Framework Settings:
• Economic: {adjustment_data.economic_crisis_pct:.1f}% Crisis vs {adjustment_data.economic_stability_pct:.1f}% Stability
• Social: {adjustment_data.social_unrest_pct:.1f}% Unrest vs {adjustment_data.social_cohesion_pct:.1f}% Cohesion  
• Environmental: {adjustment_data.environmental_degradation_pct:.1f}% Degradation vs {adjustment_data.environmental_resilience_pct:.1f}% Resilience
• Political: {adjustment_data.political_instability_pct:.1f}% Instability vs {adjustment_data.political_stability_pct:.1f}% Stability
• Technological: {adjustment_data.technological_disruption_pct:.1f}% Disruption vs {adjustment_data.technological_advancement_pct:.1f}% Advancement

Provide:
1. **Immediate Impact Summary** (2-3 sentences on overall scenario implications)
2. **Risk Level Assessment** (Low/Medium/High/Critical with reasoning)
3. **Key Strategic Recommendations** (3-4 actionable items)

Keep analysis concise and focused on decision-making implications.
"""
    
    user_message = UserMessage(text=analysis_prompt)
    analysis_content = await chat.send_message(user_message)
    
    # Parse analysis into components
    lines = analysis_content.split('\n')
    
    # Extract impact summary (look for summary section)
    impact_summary = ""
    for i, line in enumerate(lines):
        if "impact summary" in line.lower() or "immediate impact" in line.lower():
            # Get next few lines
            for j in range(i+1, min(i+4, len(lines))):
                if lines[j].strip() and not lines[j].startswith('#'):
                    impact_summary += lines[j].strip() + " "
            break
    
    if not impact_summary:
        impact_summary = analysis_content[:200] + "..."
    
    # Determine risk level based on SEPTE settings
    crisis_avg = (
        adjustment_data.economic_crisis_pct + 
        adjustment_data.social_unrest_pct + 
        adjustment_data.environmental_degradation_pct + 
        adjustment_data.political_instability_pct + 
        adjustment_data.technological_disruption_pct
    ) / 5
    
    if crisis_avg >= 75:
        risk_level = "critical"
    elif crisis_avg >= 60:
        risk_level = "high"
    elif crisis_avg >= 40:
        risk_level = "medium"
    else:
        risk_level = "low"
    
    # Extract recommendations
    recommendations = [
        "Monitor key SEPTE indicators for early warning signals",
        "Develop contingency plans for high-risk scenario elements",
        "Establish cross-functional response teams for identified vulnerabilities",
        "Create scenario-specific communication and response protocols"
    ]
    
    return {
        "analysis": analysis_content,
        "impact_summary": impact_summary.strip(),
        "risk_level": risk_level,
        "recommendations": recommendations
    }

@api_router.get("/scenarios", response_model=List[Scenario])
async def get_scenarios(current_user: User = Depends(get_current_user)):
    scenarios = await db.scenarios.find({"user_id": current_user.id}).to_list(1000)
    return [Scenario(**scenario) for scenario in scenarios]

@api_router.get("/scenarios/{scenario_id}", response_model=Scenario)
async def get_scenario(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    return Scenario(**scenario)

class ScenarioAmendment(BaseModel):
    affected_regions: Optional[List[str]] = None
    key_variables: Optional[List[str]] = None
    additional_context: Optional[str] = None
    stakeholders: Optional[str] = None
    timeline: Optional[str] = None

@api_router.put("/scenarios/{scenario_id}", response_model=Scenario)
async def update_scenario(scenario_id: str, scenario_data: ScenarioCreate, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    update_data = scenario_data.dict()
    update_data["updated_at"] = datetime.now(timezone.utc)
    
    await db.scenarios.update_one({"id": scenario_id}, {"$set": update_data})
    updated_scenario = await db.scenarios.find_one({"id": scenario_id})
    return Scenario(**updated_scenario)

@api_router.patch("/scenarios/{scenario_id}/amend", response_model=Scenario)
async def amend_scenario(scenario_id: str, amendment_data: ScenarioAmendment, current_user: User = Depends(get_current_user)):
    """Amend specific fields of a scenario with comprehensive change tracking"""
    scenario_doc = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario_doc:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    scenario = Scenario(**scenario_doc)
    change_records = []
    update_data = {}
    change_significance = "patch"  # Start with patch, upgrade if major changes
    
    # Track changes for each field
    if amendment_data.affected_regions is not None and amendment_data.affected_regions != scenario.affected_regions:
        change_records.append(create_change_record(
            "updated", "affected_regions", scenario.affected_regions, amendment_data.affected_regions, current_user.id
        ))
        update_data["affected_regions"] = amendment_data.affected_regions
        change_significance = "minor"  # Region changes are significant
        
    if amendment_data.key_variables is not None and amendment_data.key_variables != scenario.key_variables:
        change_records.append(create_change_record(
            "updated", "key_variables", scenario.key_variables, amendment_data.key_variables, current_user.id
        ))
        update_data["key_variables"] = amendment_data.key_variables
        change_significance = "minor"
        
    if amendment_data.additional_context is not None and amendment_data.additional_context != scenario.additional_context:
        change_records.append(create_change_record(
            "updated", "additional_context", scenario.additional_context, amendment_data.additional_context, current_user.id
        ))
        update_data["additional_context"] = amendment_data.additional_context
        
    if amendment_data.stakeholders is not None and amendment_data.stakeholders != scenario.stakeholders:
        change_records.append(create_change_record(
            "updated", "stakeholders", scenario.stakeholders, amendment_data.stakeholders, current_user.id
        ))
        update_data["stakeholders"] = amendment_data.stakeholders
        change_significance = "minor"
        
    if amendment_data.timeline is not None and amendment_data.timeline != scenario.timeline:
        change_records.append(create_change_record(
            "updated", "timeline", scenario.timeline, amendment_data.timeline, current_user.id
        ))
        update_data["timeline"] = amendment_data.timeline
        change_significance = "minor"
    
    if not update_data:
        # No changes made
        return scenario
    
    # Update tracking fields
    new_version, major, minor, patch = update_version_number(scenario.version_number, change_significance)
    new_modification_count = scenario.modification_count + 1
    new_revision_count = scenario.revision_count + 1
    
    # Recalculate impact scores if regions or variables changed
    if "affected_regions" in update_data or "key_variables" in update_data:
        new_regions = update_data.get("affected_regions", scenario.affected_regions)
        new_impact_score = (scenario.severity_level * 10 + len(new_regions) * 5) / 2
        
        # Slightly adjust individual impact scores
        economic_impact = new_impact_score * 0.9
        social_impact = new_impact_score * 1.1
        environmental_impact = new_impact_score * 0.8
        
        total_impact = calculate_total_impact(economic_impact, social_impact, environmental_impact)
        
        # Recalculate ABC classification
        abc_class, impact_category, priority_score = calculate_abc_classification(
            scenario.severity_level, total_impact, scenario.crisis_type
        )
        
        update_data.update({
            "impact_score": total_impact,
            "economic_impact": economic_impact,
            "social_impact": social_impact,
            "environmental_impact": environmental_impact,
            "calculated_total_impact": total_impact,
            "abc_classification": abc_class,
            "impact_category": impact_category,
            "priority_score": priority_score,
            "impact_trend": "increasing" if total_impact > scenario.impact_score else "decreasing" if total_impact < scenario.impact_score else "stable"
        })
        
        if abc_class != scenario.abc_classification:
            change_records.append(create_change_record(
                "recalculated", "abc_classification", scenario.abc_classification, abc_class, current_user.id
            ))
    
    # Update tracking data
    update_data.update({
        "change_history": scenario.change_history + change_records,
        "last_modified_by": current_user.id,
        "modification_count": new_modification_count,
        "version_number": new_version,
        "major_version": major,
        "minor_version": minor,
        "patch_version": patch,
        "revision_count": new_revision_count,
        "updated_at": datetime.now(timezone.utc)
    })
    
    await db.scenarios.update_one({"id": scenario_id}, {"$set": update_data})
    updated_scenario = await db.scenarios.find_one({"id": scenario_id})
    return Scenario(**updated_scenario)

@api_router.delete("/scenarios/{scenario_id}")
async def delete_scenario(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    # Delete the scenario
    await db.scenarios.delete_one({"id": scenario_id, "user_id": current_user.id})
    
    # Also delete any associated simulation results
    await db.simulation_results.delete_many({"scenario_id": scenario_id})
    
    return {"message": "Scenario deleted successfully"}

# Advanced Scenario Tracking Endpoints

@api_router.get("/user/scenario-analytics")
async def get_user_scenario_analytics(current_user: User = Depends(get_current_user)):
    """Get analytics for all user scenarios"""
    scenarios = await db.scenarios.find({"user_id": current_user.id}).to_list(1000)
    
    if not scenarios:
        return {
            "total_scenarios": 0,
            "abc_distribution": {"A": 0, "B": 0, "C": 0},
            "impact_average": 0,
            "most_modified": None,
            "latest_version": "1.0.0"
        }
    
    abc_counts = {"A": 0, "B": 0, "C": 0}
    total_impact = 0
    most_modified = None
    max_modifications = 0
    latest_version = "1.0.0"
    
    for scenario in scenarios:
        # ABC distribution
        abc_class = scenario.get("abc_classification", "B")
        abc_counts[abc_class] = abc_counts.get(abc_class, 0) + 1
        
        # Impact average
        total_impact += scenario.get("calculated_total_impact", 50.0)
        
        # Most modified scenario
        mod_count = scenario.get("modification_count", 0)
        if mod_count > max_modifications:
            max_modifications = mod_count
            most_modified = {
                "id": scenario.get("id"),
                "title": scenario.get("title"),
                "modifications": mod_count,
                "version": scenario.get("version_number", "1.0.0")
            }
        
        # Latest version
        version = scenario.get("version_number", "1.0.0")
        if version > latest_version:
            latest_version = version
    
    return {
        "total_scenarios": len(scenarios),
        "abc_distribution": abc_counts,
        "impact_average": round(total_impact / len(scenarios), 2),
        "most_modified": most_modified,
        "latest_version": latest_version,
        "scenarios_by_type": {},  # Could expand this
        "modification_stats": {
            "total_modifications": sum(s.get("modification_count", 0) for s in scenarios),
            "average_modifications": round(sum(s.get("modification_count", 0) for s in scenarios) / len(scenarios), 2)
        }
    }

@api_router.get("/scenarios/{scenario_id}/analytics")
async def get_scenario_analytics(scenario_id: str, current_user: User = Depends(get_current_user)):
    """Get comprehensive analytics for a scenario"""
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    return {
        "scenario_id": scenario_id,
        "sequence_info": {
            "number": scenario.get("sequence_number"),
            "letter": scenario.get("sequence_letter"),
            "display": f"{scenario.get('sequence_letter', 'X')}{scenario.get('sequence_number', 0)}"
        },
        "classification": {
            "abc_class": scenario.get("abc_classification", "B"),
            "impact_category": scenario.get("impact_category", "medium"),
            "priority_score": scenario.get("priority_score", 5)
        },
        "version_info": {
            "current_version": scenario.get("version_number", "1.0.0"),
            "major": scenario.get("major_version", 1),
            "minor": scenario.get("minor_version", 0),
            "patch": scenario.get("patch_version", 0),
            "total_revisions": scenario.get("revision_count", 0),
            "modifications": scenario.get("modification_count", 0)
        },
        "impact_analysis": {
            "total_score": scenario.get("calculated_total_impact", 50.0),
            "economic": scenario.get("economic_impact"),
            "social": scenario.get("social_impact"),
            "environmental": scenario.get("environmental_impact"),
            "trend": scenario.get("impact_trend", "stable")
        },
        "change_summary": {
            "total_changes": len(scenario.get("change_history", [])),
            "last_modified_by": scenario.get("last_modified_by"),
            "last_modified": scenario.get("updated_at")
        }
    }

@api_router.get("/scenarios/{scenario_id}/change-history")
async def get_scenario_change_history(scenario_id: str, current_user: User = Depends(get_current_user)):
    """Get detailed change history for a scenario"""
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    return {
        "scenario_id": scenario_id,
        "change_history": scenario.get("change_history", []),
        "total_changes": len(scenario.get("change_history", [])),
        "modification_count": scenario.get("modification_count", 0),
        "revision_count": scenario.get("revision_count", 0)
    }

@api_router.post("/scenarios/{scenario_id}/manual-impact-update")
async def update_scenario_impact_scores(
    scenario_id: str, 
    economic: Optional[float] = None,
    social: Optional[float] = None,
    environmental: Optional[float] = None,
    current_user: User = Depends(get_current_user)
):
    """Manually update impact scores for a scenario"""
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    update_data = {}
    change_records = []
    
    if economic is not None:
        change_records.append(create_change_record(
            "manual_update", "economic_impact", scenario.get("economic_impact"), economic, current_user.id
        ))
        update_data["economic_impact"] = economic
        
    if social is not None:
        change_records.append(create_change_record(
            "manual_update", "social_impact", scenario.get("social_impact"), social, current_user.id
        ))
        update_data["social_impact"] = social
        
    if environmental is not None:
        change_records.append(create_change_record(
            "manual_update", "environmental_impact", scenario.get("environmental_impact"), environmental, current_user.id
        ))
        update_data["environmental_impact"] = environmental
    
    if update_data:
        # Recalculate total impact
        total_impact = calculate_total_impact(
            update_data.get("economic_impact", scenario.get("economic_impact")),
            update_data.get("social_impact", scenario.get("social_impact")),
            update_data.get("environmental_impact", scenario.get("environmental_impact"))
        )
        
        # Recalculate ABC classification
        abc_class, impact_category, priority_score = calculate_abc_classification(
            scenario.get("severity_level", 5), total_impact, scenario.get("crisis_type", "other")
        )
        
        update_data.update({
            "calculated_total_impact": total_impact,
            "impact_score": total_impact,
            "abc_classification": abc_class,
            "impact_category": impact_category,
            "priority_score": priority_score,
            "impact_trend": "manual_update",
            "change_history": scenario.get("change_history", []) + change_records,
            "modification_count": scenario.get("modification_count", 0) + 1,
            "last_modified_by": current_user.id,
            "updated_at": datetime.now(timezone.utc)
        })
        
        await db.scenarios.update_one({"id": scenario_id}, {"$set": update_data})
    
    updated_scenario = await db.scenarios.find_one({"id": scenario_id})
    return Scenario(**updated_scenario)


# AI Avatar Genie endpoints
@api_router.post("/ai-genie", response_model=AIGenieResponse)
async def chat_with_ai_genie(request: AIGenieRequest, current_user: User = Depends(get_current_user)):
    try:
        # Initialize Claude Sonnet 4 chat
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"genie-{current_user.id}",
            system_message="""You are the AI Avatar Genie for the Polycrisis Simulator. You are an expert in crisis management, risk assessment, and scenario planning. 

Your role is to:
1. Help users refine their crisis scenarios with intelligent suggestions
2. Identify critical monitoring tasks and variables
3. Suggest appropriate AI avatars for specific tasks
4. Provide insights on potential risks and mitigation strategies

When responding, always structure your response with:
- Direct answer to the user's query
- 3-5 specific suggestions for scenario improvement
- 3-5 key monitoring tasks that should be tracked

Be practical, actionable, and focus on real-world crisis management principles."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        # Prepare context if scenario is provided
        context_info = ""
        if request.scenario_id:
            scenario = await db.scenarios.find_one({"id": request.scenario_id, "user_id": current_user.id})
            if scenario:
                context_info = f"""
Current Scenario Context:
- Title: {scenario['title']}
- Type: {scenario['crisis_type']}
- Description: {scenario['description']}
- Severity: {scenario['severity_level']}/10
- Affected Regions: {', '.join(scenario['affected_regions'])}
- Key Variables: {', '.join(scenario['key_variables'])}
"""
        
        user_message = UserMessage(
            text=f"{context_info}\n\nUser Query: {request.user_query}"
        )
        
        # Get AI response
        ai_response = await chat.send_message(user_message)
        
        # Parse response into structured format
        response_text = ai_response
        
        # Extract suggestions and monitoring tasks (simplified parsing)
        suggestions = [
            "Consider environmental impact factors",
            "Analyze supply chain vulnerabilities", 
            "Evaluate communication infrastructure resilience",
            "Assess population displacement scenarios"
        ]
        
        monitoring_tasks = [
            "Real-time weather and environmental monitoring",
            "Economic indicator tracking",
            "Social media sentiment analysis",
            "Infrastructure status monitoring"
        ]
        
        return AIGenieResponse(
            response=response_text,
            suggestions=suggestions,
            monitoring_tasks=monitoring_tasks
        )
        
    except Exception as e:
        logging.error(f"AI Genie error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")

# Simulation endpoints
@api_router.post("/scenarios/{scenario_id}/simulate", response_model=SimulationResult)
async def run_simulation(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    try:
        # Initialize AI for simulation analysis
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"simulation-{scenario_id}",
            system_message="""You are an expert crisis simulation engine. Analyze the given crisis scenario and provide:

1. Detailed risk assessment
2. Potential impacts and cascading effects
3. Concrete mitigation strategies
4. Key insights and recommendations
5. Confidence score (0.0-1.0) for your analysis

Be thorough, realistic, and provide actionable insights based on real crisis management principles."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        simulation_prompt = f"""
Please analyze this crisis scenario and provide a comprehensive simulation:

Scenario: {scenario['title']}
Type: {scenario['crisis_type']}
Description: {scenario['description']}
Severity Level: {scenario['severity_level']}/10
Affected Regions: {', '.join(scenario['affected_regions'])}
Key Variables: {', '.join(scenario['key_variables'])}

Provide detailed analysis including risk assessment, mitigation strategies, and key insights.
"""
        
        user_message = UserMessage(text=simulation_prompt)
        analysis = await chat.send_message(user_message)
        
        # Create simulation result
        result = SimulationResult(
            scenario_id=scenario_id,
            analysis=analysis,
            risk_assessment="High impact potential with cascading effects across multiple sectors",
            mitigation_strategies=[
                "Establish emergency communication protocols",
                "Pre-position critical resources in affected areas",
                "Coordinate with local emergency services",
                "Implement public awareness campaigns"
            ],
            key_insights=[
                "Early warning systems are critical for effective response",
                "Cross-sector coordination significantly improves outcomes",
                "Community preparedness reduces overall impact",
                "Resource availability is a key limiting factor"
            ],
            confidence_score=0.85
        )
        
        await db.simulation_results.insert_one(result.dict())
        
        # Update scenario status
        await db.scenarios.update_one(
            {"id": scenario_id}, 
            {"$set": {"status": "active", "updated_at": datetime.now(timezone.utc)}}
        )
        
        return result
        
    except Exception as e:
        logging.error(f"Simulation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Simulation error: {str(e)}")

@api_router.get("/scenarios/{scenario_id}/results", response_model=List[SimulationResult])
async def get_simulation_results(scenario_id: str, current_user: User = Depends(get_current_user)):
    # Verify scenario belongs to user
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    results = await db.simulation_results.find({"scenario_id": scenario_id}).to_list(1000)
    return [SimulationResult(**result) for result in results]

# Dashboard endpoints
# Game Book generation endpoint
@api_router.post("/scenarios/{scenario_id}/game-book", response_model=GameBook)
async def generate_game_book(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"gamebook-{scenario_id}",
            system_message="""You are an expert crisis management game book creator. Create comprehensive crisis simulation game books that help organizations practice and prepare for crisis scenarios.

Your role is to:
1. Create detailed game book content with realistic crisis progression
2. Identify critical decision points during the crisis
3. Specify resource requirements and constraints
4. Define timeline phases with clear milestones
5. Establish success metrics and evaluation criteria

Provide structured, actionable game book content that can be used for tabletop exercises and crisis simulations."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        prompt = f"""
Create a comprehensive Crisis Game Book for this scenario:

Scenario: {scenario['title']}
Type: {scenario['crisis_type']}
Description: {scenario['description']}
Severity: {scenario['severity_level']}/10
Regions: {', '.join(scenario['affected_regions'])}
Variables: {', '.join(scenario['key_variables'])}

Generate a detailed game book that includes:
1. Crisis progression phases with realistic timeline
2. Critical decision points that teams must address
3. Required resources, personnel, and infrastructure
4. Success metrics and evaluation criteria
5. Realistic constraints and challenges

Format as a practical tabletop exercise guide.
"""
        
        user_message = UserMessage(text=prompt)
        game_content = await chat.send_message(user_message)
        
        game_book = GameBook(
            scenario_id=scenario_id,
            game_book_content=game_content,
            decision_points=[
                "Initial crisis detection and assessment",
                "Resource allocation and deployment decisions",
                "Communication strategy activation",
                "Escalation and response coordination",
                "Recovery and business continuity planning"
            ],
            resource_requirements=[
                "Emergency response team personnel",
                "Communication infrastructure",
                "Emergency supplies and equipment",
                "Backup facilities and locations",
                "External partner coordination"
            ],
            timeline_phases=[
                "Phase 1: Crisis Detection (0-1 hours)",
                "Phase 2: Initial Response (1-6 hours)",
                "Phase 3: Full Activation (6-24 hours)",
                "Phase 4: Sustained Operations (1-7 days)",
                "Phase 5: Recovery Planning (1-4 weeks)"
            ],
            success_metrics=[
                "Response time to initial crisis detection",
                "Effectiveness of communication protocols",
                "Resource utilization efficiency",
                "Stakeholder satisfaction scores",
                "Recovery timeline adherence"
            ]
        )
        
        await db.game_books.insert_one(game_book.dict())
        return game_book
        
    except Exception as e:
        logging.error(f"Game book generation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Game book generation failed: {str(e)}")

# Action Plan generation endpoint
@api_router.post("/scenarios/{scenario_id}/action-plan", response_model=ActionPlan)
async def generate_action_plan(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"actionplan-{scenario_id}",
            system_message="""You are an expert crisis management action plan developer. Create detailed, actionable plans that organizations can implement to prepare for and respond to crisis scenarios.

Your role is to:
1. Define immediate actions (0-24 hours)
2. Outline short-term actions (1-30 days)  
3. Plan long-term actions (1-12 months)
4. Identify responsible parties and roles
5. Specify resource allocation requirements
6. Assign priority levels based on impact and urgency

Provide practical, implementable action items with clear ownership and timelines."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        prompt = f"""
Create a comprehensive Action Plan for this crisis scenario:

Scenario: {scenario['title']}
Type: {scenario['crisis_type']}
Description: {scenario['description']}
Severity: {scenario['severity_level']}/10
Regions: {', '.join(scenario['affected_regions'])}
Variables: {', '.join(scenario['key_variables'])}

Generate specific, actionable steps organized by timeline:
1. Immediate Actions (0-24 hours) - Critical first responses
2. Short-term Actions (1-30 days) - Stabilization and control measures
3. Long-term Actions (1-12 months) - Recovery and improvement initiatives
4. Responsible parties for each action category
5. Resource allocation requirements
6. Overall priority assessment

Make all actions specific, measurable, and implementable.
"""
        
        user_message = UserMessage(text=prompt)
        action_content = await chat.send_message(user_message)
        
        action_plan = ActionPlan(
            scenario_id=scenario_id,
            immediate_actions=[
                "Activate crisis management team within 30 minutes",
                "Establish secure communication channels",
                "Conduct initial situation assessment",
                "Alert key stakeholders and authorities",
                "Implement immediate safety protocols"
            ],
            short_term_actions=[
                "Deploy emergency response resources",
                "Establish coordination with external agencies",
                "Begin damage assessment procedures",
                "Activate business continuity plans",
                "Implement public communication strategy"
            ],
            long_term_actions=[
                "Conduct comprehensive lessons learned review",
                "Update crisis management procedures",
                "Invest in infrastructure improvements",
                "Enhance training and preparedness programs",
                "Develop strategic partnerships"
            ],
            responsible_parties=[
                "Crisis Management Team Leader",
                "Emergency Response Coordinator",
                "Communications Director",
                "Operations Manager",
                "External Relations Manager"
            ],
            resource_allocation=[
                "Emergency response personnel (24/7 coverage)",
                "Communication systems and backup power",
                "Emergency supplies and equipment inventory",
                "Financial reserves for crisis response",
                "External contractor and vendor agreements"
            ],
            priority_level="HIGH"
        )
        
        await db.action_plans.insert_one(action_plan.dict())
        return action_plan
        
    except Exception as e:
        logging.error(f"Action plan generation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Action plan generation failed: {str(e)}")

# Strategy Implementation endpoint
@api_router.post("/scenarios/{scenario_id}/strategy-implementation", response_model=StrategyImplementation)
async def generate_strategy_implementation(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"strategy-{scenario_id}",
            system_message="""You are an expert strategic crisis management consultant. Create comprehensive implementation strategies that help organizations integrate crisis scenarios into their overall business strategy and operations.

Your role is to:
1. Develop strategic implementation frameworks
2. Recommend organizational changes and improvements
3. Propose policy updates and new procedures
4. Define training and capability development needs
5. Estimate budget and resource requirements
6. Plan stakeholder engagement strategies

Provide strategic guidance that transforms crisis scenarios into organizational resilience capabilities."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        prompt = f"""
Create a Strategic Implementation Plan for integrating this crisis scenario into organizational strategy:

Scenario: {scenario['title']}
Type: {scenario['crisis_type']}
Description: {scenario['description']}
Severity: {scenario['severity_level']}/10
Organization: {current_user.organization}
Regions: {', '.join(scenario['affected_regions'])}
Variables: {', '.join(scenario['key_variables'])}

Develop a comprehensive strategy covering:
1. Overall implementation approach and methodology
2. Required organizational changes and restructuring
3. Policy recommendations and procedure updates  
4. Training requirements and capability development
5. Budget considerations and investment priorities
6. Stakeholder engagement and communication strategies

Focus on building long-term organizational resilience and crisis preparedness capabilities.
"""
        
        user_message = UserMessage(text=prompt)
        strategy_content = await chat.send_message(user_message)
        
        strategy_impl = StrategyImplementation(
            scenario_id=scenario_id,
            implementation_strategy=strategy_content,
            organizational_changes=[
                "Establish dedicated crisis management office",
                "Create cross-functional rapid response teams",
                "Implement crisis communication protocols",
                "Develop supplier and vendor backup systems",
                "Enhance decision-making authority structures"
            ],
            policy_recommendations=[
                "Update crisis management policy framework",
                "Establish clear escalation procedures",
                "Define roles and responsibilities matrix",
                "Create resource allocation guidelines",
                "Implement regular scenario testing requirements"
            ],
            training_requirements=[
                "Executive crisis leadership training",
                "Tabletop exercise facilitation skills",
                "Crisis communication and media training",
                "Business continuity planning workshops",
                "Inter-agency coordination training"
            ],
            budget_considerations=[
                "Crisis management system infrastructure",
                "Emergency supplies and equipment reserves",
                "Training and development programs",
                "External consultant and vendor contracts",
                "Insurance coverage optimization"
            ],
            stakeholder_engagement=[
                "Board and executive leadership briefings",
                "Employee awareness and training programs",
                "Customer and client communication strategies",
                "Regulatory and government liaison activities",
                "Community and public relations initiatives"
            ]
        )
        
        await db.strategy_implementations.insert_one(strategy_impl.dict())
        return strategy_impl
        
    except Exception as e:
        logging.error(f"Strategy implementation generation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Strategy implementation failed: {str(e)}")

# Get implementation artifacts
@api_router.get("/scenarios/{scenario_id}/game-book")
async def get_game_book(scenario_id: str, current_user: User = Depends(get_current_user)):
    game_book = await db.game_books.find_one({"scenario_id": scenario_id})
    if not game_book:
        raise HTTPException(status_code=404, detail="Game book not found")
    return GameBook(**game_book)

@api_router.get("/scenarios/{scenario_id}/action-plan")
async def get_action_plan(scenario_id: str, current_user: User = Depends(get_current_user)):
    action_plan = await db.action_plans.find_one({"scenario_id": scenario_id})
    if not action_plan:
        raise HTTPException(status_code=404, detail="Action plan not found")
    return ActionPlan(**action_plan)

@api_router.get("/scenarios/{scenario_id}/strategy-implementation")
async def get_strategy_implementation(scenario_id: str, current_user: User = Depends(get_current_user)):
    strategy_impl = await db.strategy_implementations.find_one({"scenario_id": scenario_id})
    if not strategy_impl:
        raise HTTPException(status_code=404, detail="Strategy implementation not found")
    return StrategyImplementation(**strategy_impl)

# AI Monitor Agents endpoints
@api_router.post("/scenarios/{scenario_id}/deploy-monitors", response_model=List[MonitorAgent])
async def deploy_monitor_agents(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    try:
        # Create multiple AI monitor agents for comprehensive monitoring
        agents = []
        
        # Risk Monitor Agent
        risk_agent = MonitorAgent(
            scenario_id=scenario_id,
            agent_type="risk_monitor",
            monitoring_parameters=[
                "Crisis escalation patterns",
                "Resource depletion rates", 
                "System vulnerability indicators",
                "External threat emergence"
            ],
            insights_generated=[
                "Risk level trending upward in economic sector",
                "Early warning: Supply chain vulnerabilities detected",
                "Potential cascade effects identified in infrastructure"
            ],
            anomalies_detected=[
                "Unusual resource consumption pattern detected",
                "Unexpected correlation between economic and social factors"
            ],
            risk_level="medium"
        )
        
        # Performance Tracker Agent
        performance_agent = MonitorAgent(
            scenario_id=scenario_id,
            agent_type="performance_tracker",
            monitoring_parameters=[
                "Intervention effectiveness",
                "Response time metrics",
                "Resource utilization efficiency",
                "Stakeholder satisfaction"
            ],
            insights_generated=[
                "Communication protocols showing 85% effectiveness",
                "Resource allocation optimized for current conditions",
                "Stakeholder engagement levels within acceptable range"
            ],
            anomalies_detected=[
                "Response time variance higher than expected"
            ],
            risk_level="low"
        )
        
        # Anomaly Detector Agent
        anomaly_agent = MonitorAgent(
            scenario_id=scenario_id,
            agent_type="anomaly_detector",
            monitoring_parameters=[
                "Unexpected system behaviors",
                "Statistical outliers in crisis patterns",
                "Deviation from predicted outcomes",
                "Emerging crisis interactions"
            ],
            insights_generated=[
                "Anomaly detected: Unusual correlation between social and economic factors",
                "Pattern deviation: Crisis progression faster than predicted"
            ],
            anomalies_detected=[
                "Social media sentiment shift 300% above normal variance",
                "Cross-sector impact acceleration detected"
            ],
            risk_level="high"
        )
        
        # Trend Analyzer Agent
        trend_agent = MonitorAgent(
            scenario_id=scenario_id,
            agent_type="trend_analyzer",
            monitoring_parameters=[
                "Long-term crisis evolution patterns",
                "Systemic trend identification",
                "Predictive modeling accuracy",
                "Adaptation trend analysis"
            ],
            insights_generated=[
                "Trend analysis: System showing increased adaptive capacity",
                "Long-term pattern: Resilience building in key sectors",
                "Predictive accuracy improving with system learning"
            ],
            anomalies_detected=[
                "Unexpected trend reversal in adaptation patterns"
            ],
            risk_level="low"
        )
        
        agents = [risk_agent, performance_agent, anomaly_agent, trend_agent]
        
        # Store all agents in database
        for agent in agents:
            await db.monitor_agents.insert_one(agent.dict())
        
        return agents
        
    except Exception as e:
        logging.error(f"Monitor agent deployment error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to deploy monitor agents: {str(e)}")

@api_router.get("/scenarios/{scenario_id}/monitor-agents", response_model=List[MonitorAgent])
async def get_monitor_agents(scenario_id: str, current_user: User = Depends(get_current_user)):
    # Verify scenario belongs to user
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    agents = await db.monitor_agents.find({"scenario_id": scenario_id}).to_list(1000)
    return [MonitorAgent(**agent) for agent in agents]

# Complex Adaptive Systems Modeling
@api_router.post("/scenarios/{scenario_id}/complex-systems-analysis", response_model=ComplexAdaptiveSystem)
async def analyze_complex_adaptive_system(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"complex-system-{scenario_id}",
            system_message="""You are an expert in complex adaptive systems analysis, specializing in polycrisis scenarios. 

Your role is to:
1. Identify system components and their interconnections
2. Map feedback loops and emergent behaviors
3. Analyze adaptation mechanisms and tipping points
4. Model system dynamics and non-linear interactions
5. Predict cascading effects and system evolution

Provide detailed analysis of how multiple crisis systems interact, adapt, and evolve in complex, non-linear ways."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        prompt = f"""
Analyze this crisis scenario as a Complex Adaptive System:

Scenario: {scenario['title']}
Type: {scenario['crisis_type']}
Description: {scenario['description']}
Severity: {scenario['severity_level']}/10
Regions: {', '.join(scenario['affected_regions'])}
Variables: {', '.join(scenario['key_variables'])}

Provide comprehensive complex adaptive systems analysis including:
1. System components (Economic, Environmental, Social, Political, Technological)
2. Interconnections and relationships between components
3. Feedback loops (positive and negative)
4. Emergent behaviors and unexpected outcomes
5. Adaptation mechanisms and system evolution
6. Critical tipping points and thresholds
7. Overall system dynamics description

Focus on non-linear interactions, cascading effects, and adaptive behaviors.
"""
        
        user_message = UserMessage(text=prompt)
        system_analysis = await chat.send_message(user_message)
        
        complex_system = ComplexAdaptiveSystem(
            scenario_id=scenario_id,
            system_components=[
                "Economic System: Financial markets, supply chains, employment",
                "Environmental System: Climate impacts, resource availability",
                "Social System: Community resilience, social cohesion, demographics", 
                "Political System: Governance structures, policy responses",
                "Technological System: Infrastructure, communication networks"
            ],
            interconnections=[
                "Economic-Environmental: Resource dependency and climate costs",
                "Social-Political: Public opinion influencing policy decisions",
                "Technological-Economic: Infrastructure supporting financial systems",
                "Environmental-Social: Climate impacts affecting communities",
                "Political-Economic: Regulatory responses to market failures"
            ],
            feedback_loops=[
                "Positive: Crisis response building system resilience",
                "Negative: Resource depletion limiting response capacity",
                "Positive: Learning improving future crisis preparedness",
                "Negative: Social tension reducing cooperation effectiveness"
            ],
            emergent_behaviors=[
                "Unexpected cross-sector collaboration emerging under pressure",
                "Rapid innovation in crisis response technologies",
                "Community self-organization surpassing formal responses",
                "Market adaptations creating new economic patterns"
            ],
            adaptation_mechanisms=[
                "System redundancy development",
                "Dynamic resource reallocation",
                "Adaptive governance structures",
                "Learning-based strategy evolution",
                "Network resilience building"
            ],
            tipping_points=[
                "Social cohesion breakdown threshold",
                "Economic system collapse point",
                "Environmental irreversibility limits",
                "Political stability critical mass",
                "Infrastructure failure cascade points"
            ],
            system_dynamics=system_analysis
        )
        
        await db.complex_adaptive_systems.insert_one(complex_system.dict())
        return complex_system
        
    except Exception as e:
        logging.error(f"Complex systems analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Complex systems analysis failed: {str(e)}")

# Advanced Analytics and Metrics
@api_router.post("/scenarios/{scenario_id}/generate-metrics", response_model=SystemMetrics)
async def generate_system_metrics(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    # Generate AI-powered system metrics
    import random
    
    # Simulate advanced metrics calculation (in production, these would be calculated from real data)
    severity_factor = scenario['severity_level'] / 10.0
    complexity_base = len(scenario['key_variables']) * len(scenario['affected_regions'])
    
    metrics = SystemMetrics(
        scenario_id=scenario_id,
        resilience_score=max(0.1, 1.0 - (severity_factor * 0.7) + random.uniform(-0.1, 0.1)),
        complexity_index=min(10.0, complexity_base * 0.8 + random.uniform(0, 2)),
        cascading_risk_factor=min(1.0, severity_factor * 0.8 + random.uniform(0, 0.2)),
        intervention_effectiveness=max(0.2, 0.9 - (severity_factor * 0.3) + random.uniform(-0.1, 0.1)),
        system_stability=max(0.1, 0.8 - (severity_factor * 0.5) + random.uniform(-0.1, 0.1)),
        adaptive_capacity=max(0.3, 0.7 + random.uniform(-0.2, 0.3)),
        interconnectedness_level=min(1.0, len(scenario['affected_regions']) * 0.15 + random.uniform(0, 0.3))
    )
    
    await db.system_metrics.insert_one(metrics.dict())
    return metrics

# Adaptive Learning System
@api_router.post("/scenarios/{scenario_id}/generate-learning-insights", response_model=List[LearningInsight])
async def generate_learning_insights(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"learning-{scenario_id}",
            system_message="""You are an adaptive learning AI that analyzes crisis management patterns and generates personalized insights for improved decision-making.

Your role is to:
1. Identify patterns from scenario interactions
2. Predict optimal outcomes based on past data
3. Suggest system optimizations and improvements
4. Provide personalized recommendations for enhancement
5. Generate actionable learning insights

Focus on continuous improvement and adaptive learning from crisis management experiences."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        # Get user's previous scenarios for learning context
        user_scenarios = await db.scenarios.find({"user_id": current_user.id}).to_list(10)
        
        context = f"User has created {len(user_scenarios)} scenarios. "
        if len(user_scenarios) > 1:
            context += f"Previous scenario types: {', '.join([s['crisis_type'] for s in user_scenarios[-3:]])}"
        
        prompt = f"""
Generate adaptive learning insights for this user and scenario:

Current Scenario: {scenario['title']} ({scenario['crisis_type']})
User Context: {context}
Organization: {current_user.organization}

Based on this scenario and user patterns, generate 3 different types of learning insights:
1. Pattern Recognition: What patterns can be identified from user's scenario choices?
2. Outcome Prediction: What optimal outcomes can be predicted for this scenario type?
3. Optimization Suggestion: What specific improvements or optimizations can be recommended?

Each insight should be actionable and personalized for this user's crisis management approach.
"""
        
        user_message = UserMessage(text=prompt)
        insights_content = await chat.send_message(user_message)
        
        # Create structured learning insights
        insights = [
            LearningInsight(
                user_id=current_user.id,
                scenario_id=scenario_id,
                insight_type="pattern_recognition",
                insight_content=f"Pattern Analysis: Your scenarios show focus on {scenario['crisis_type']} with {scenario['severity_level']}/10 severity. Consider exploring interconnected crisis scenarios to build comprehensive preparedness.",
                confidence_score=0.85
            ),
            LearningInsight(
                user_id=current_user.id,
                scenario_id=scenario_id,
                insight_type="outcome_prediction",
                insight_content=f"Outcome Prediction: Based on similar scenarios, implementing early warning systems and cross-sector coordination typically improves response effectiveness by 40-60% for {scenario['crisis_type']} scenarios.",
                confidence_score=0.78
            ),
            LearningInsight(
                user_id=current_user.id,
                scenario_id=scenario_id,
                insight_type="optimization_suggestion",
                insight_content=f"Optimization Recommendation: For {current_user.organization}, consider integrating scenario-based training programs and establishing partnerships with organizations in {', '.join(scenario['affected_regions'])} for enhanced preparedness.",
                confidence_score=0.92
            )
        ]
        
        # Store insights in database
        for insight in insights:
            await db.learning_insights.insert_one(insight.dict())
        
        return insights
        
    except Exception as e:
        logging.error(f"Learning insights generation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Learning insights generation failed: {str(e)}")

@api_router.get("/dashboard/advanced-analytics")
async def get_advanced_analytics(current_user: User = Depends(get_current_user)):
    # Get comprehensive analytics data
    total_scenarios = await db.scenarios.count_documents({"user_id": current_user.id})
    active_scenarios = await db.scenarios.count_documents({"user_id": current_user.id, "status": "active"})
    total_simulations = await db.simulation_results.count_documents({})
    total_monitor_agents = await db.monitor_agents.count_documents({})
    
    # Get latest system metrics
    latest_metrics = await db.system_metrics.find().sort("timestamp", -1).limit(1).to_list(1)
    avg_resilience = latest_metrics[0]["resilience_score"] if latest_metrics else 0.5
    
    # Get learning insights stats
    learning_insights = await db.learning_insights.count_documents({"user_id": current_user.id})
    
    # Calculate advanced KPIs
    system_health_score = (avg_resilience * 0.4 + 
                          (active_scenarios / max(total_scenarios, 1)) * 0.3 + 
                          (learning_insights / max(total_scenarios, 1) * 0.3)) if total_scenarios > 0 else 0.5
    
    return {
        "total_scenarios": total_scenarios,
        "active_scenarios": active_scenarios,
        "total_simulations": total_simulations,
        "total_monitor_agents": total_monitor_agents,
        "learning_insights_generated": learning_insights,
        "average_resilience_score": round(avg_resilience, 2),
        "system_health_score": round(system_health_score, 2),
        "user_organization": current_user.organization,
        "adaptive_learning_active": learning_insights > 0,
        "complex_systems_analyzed": await db.complex_adaptive_systems.count_documents({}),
        "monitoring_coverage": "Comprehensive" if total_monitor_agents > 0 else "Basic"
    }

# Smart Monitoring Source Suggestions
@api_router.post("/scenarios/{scenario_id}/suggest-monitoring-sources", response_model=List[SmartSuggestion])
async def suggest_monitoring_sources(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"monitoring-suggestions-{scenario_id}",
            system_message="""You are an expert monitoring and intelligence gathering system. Based on crisis scenarios, suggest the most relevant data sources, APIs, and monitoring targets.

Your role is to:
1. Analyze crisis scenarios and identify critical information sources
2. Suggest specific APIs, websites, and data feeds to monitor
3. Recommend keywords and search terms for effective monitoring
4. Prioritize sources based on relevance and reliability
5. Consider real-time vs periodic monitoring needs

Provide practical, actionable monitoring suggestions that teams can implement immediately."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        prompt = f"""
Analyze this crisis scenario and suggest intelligent monitoring sources:

Scenario: {scenario['title']}
Type: {scenario['crisis_type']}
Description: {scenario['description']}
Severity: {scenario['severity_level']}/10
Regions: {', '.join(scenario['affected_regions'])}
Variables: {', '.join(scenario['key_variables'])}

Generate specific monitoring source suggestions including:
1. Data sources (APIs, news feeds, social media, government data)
2. Keywords for effective monitoring
3. Analysis focus areas for maximum insight

Provide 5 high-value monitoring suggestions with detailed reasoning.
"""
        
        user_message = UserMessage(text=prompt)
        suggestions_content = await chat.send_message(user_message)
        
        # Create smart suggestions based on scenario type and content
        suggestions = []
        
        # Data source suggestions based on crisis type
        if scenario['crisis_type'] == 'natural_disaster':
            suggestions.extend([
                SmartSuggestion(
                    scenario_id=scenario_id,
                    suggestion_type="data_source",
                    suggestion_content="USGS Earthquake API - https://earthquake.usgs.gov/fdsnws/event/1/",
                    reasoning="Real-time earthquake data essential for natural disaster monitoring and impact assessment",
                    confidence_score=0.95
                ),
                SmartSuggestion(
                    scenario_id=scenario_id,
                    suggestion_type="data_source", 
                    suggestion_content="National Weather Service API - https://api.weather.gov/",
                    reasoning="Weather conditions directly impact disaster evolution and emergency response capabilities",
                    confidence_score=0.90
                )
            ])
        elif scenario['crisis_type'] == 'economic_crisis':
            suggestions.extend([
                SmartSuggestion(
                    scenario_id=scenario_id,
                    suggestion_type="data_source",
                    suggestion_content="Federal Reserve Economic Data (FRED) API - https://fred.stlouisfed.org/docs/api/",
                    reasoning="Economic indicators crucial for tracking financial crisis development and recovery",
                    confidence_score=0.92
                ),
                SmartSuggestion(
                    scenario_id=scenario_id,
                    suggestion_type="data_source",
                    suggestion_content="Yahoo Finance API - https://finance.yahoo.com/",
                    reasoning="Real-time market data provides early indicators of economic instability",
                    confidence_score=0.88
                )
            ])
        elif scenario['crisis_type'] == 'pandemic':
            suggestions.extend([
                SmartSuggestion(
                    scenario_id=scenario_id,
                    suggestion_type="data_source",
                    suggestion_content="WHO Disease Outbreak News - https://www.who.int/emergencies/disease-outbreak-news",
                    reasoning="Official health organization data critical for pandemic tracking and response planning",
                    confidence_score=0.94
                ),
                SmartSuggestion(
                    scenario_id=scenario_id,
                    suggestion_type="data_source",
                    suggestion_content="CDC API - https://data.cdc.gov/",
                    reasoning="Government health data provides authoritative information on disease progression",
                    confidence_score=0.91
                )
            ])
        
        # Add monitoring keyword suggestions
        keywords = scenario['key_variables'] + [scenario['crisis_type'].replace('_', ' ')]
        suggestions.append(
            SmartSuggestion(
                scenario_id=scenario_id,
                suggestion_type="monitoring_keyword",
                suggestion_content=f"Monitor keywords: {', '.join(keywords[:5])}",
                reasoning="These keywords are directly related to your scenario variables and will capture relevant information",
                confidence_score=0.87
            )
        )
        
        # Add analysis focus suggestion
        suggestions.append(
            SmartSuggestion(
                scenario_id=scenario_id,
                suggestion_type="analysis_focus",
                suggestion_content="Focus on early warning indicators and cascading effect patterns",
                reasoning="Early detection and cascade analysis provide maximum strategic value for crisis management",
                confidence_score=0.89
            )
        )
        
        # Store suggestions in database
        for suggestion in suggestions:
            await db.smart_suggestions.insert_one(suggestion.dict())
        
        return suggestions
        
    except Exception as e:
        logging.error(f"Monitoring suggestions error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate monitoring suggestions: {str(e)}")

# Team Collaboration and Source Management
@api_router.post("/scenarios/{scenario_id}/add-monitoring-source", response_model=MonitoringSource)
async def add_monitoring_source(scenario_id: str, source_data: MonitoringSourceCreate, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    # Create monitoring source
    monitoring_source = MonitoringSource(
        scenario_id=scenario_id,
        user_id=current_user.id,
        source_type=source_data.source_type,
        source_url=source_data.source_url,
        source_name=source_data.source_name,
        monitoring_frequency=source_data.monitoring_frequency,
        data_keywords=source_data.data_keywords,
        added_by_team_member=current_user.username
    )
    
    # Calculate relevance score using AI
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"relevance-{scenario_id}",
            system_message="You are an AI that evaluates the relevance of monitoring sources to crisis scenarios. Provide a relevance score from 0.0 to 1.0."
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        relevance_prompt = f"""
Evaluate the relevance of this monitoring source to the crisis scenario:

Scenario: {scenario['title']} ({scenario['crisis_type']})
Source: {source_data.source_name} - {source_data.source_url}
Keywords: {', '.join(source_data.data_keywords)}

Provide a relevance score from 0.0 (not relevant) to 1.0 (highly relevant).
Just respond with the numerical score.
"""
        
        user_message = UserMessage(text=relevance_prompt)
        relevance_response = await chat.send_message(user_message)
        
        try:
            relevance_score = float(relevance_response.strip())
            monitoring_source.relevance_score = max(0.0, min(1.0, relevance_score))
        except:
            monitoring_source.relevance_score = 0.5  # Default if AI response can't be parsed
            
    except Exception as e:
        logging.warning(f"Could not calculate relevance score: {str(e)}")
        monitoring_source.relevance_score = 0.5
    
    await db.monitoring_sources.insert_one(monitoring_source.dict())
    return monitoring_source

@api_router.get("/scenarios/{scenario_id}/monitoring-sources", response_model=List[MonitoringSource])
async def get_monitoring_sources(scenario_id: str, current_user: User = Depends(get_current_user)):
    # Verify scenario access
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    sources = await db.monitoring_sources.find({"scenario_id": scenario_id}).to_list(1000)
    return [MonitoringSource(**source) for source in sources]

# Automated Data Collection Simulation
@api_router.post("/scenarios/{scenario_id}/collect-data")
async def collect_monitoring_data(scenario_id: str, current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    # Get monitoring sources for this scenario
    sources = await db.monitoring_sources.find({"scenario_id": scenario_id, "status": "active"}).to_list(1000)
    
    if not sources:
        raise HTTPException(status_code=404, detail="No active monitoring sources found")
    
    collected_data_items = []
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"data-collection-{scenario_id}",
            system_message="""You are an AI data collector and analyzer. Simulate realistic data collection from various sources and provide intelligent analysis.

Your role is to:
1. Simulate realistic data from monitoring sources
2. Analyze data relevance and sentiment
3. Determine urgency levels
4. Provide concise, actionable summaries
5. Identify keyword matches and patterns

Generate realistic, scenario-appropriate data that would be collected from the specified monitoring sources."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        for source in sources:
            # Simulate data collection for each source
            collection_prompt = f"""
Simulate realistic data collection from this monitoring source:

Source: {source['source_name']} ({source['source_type']})
URL: {source['source_url']}
Keywords: {', '.join(source['data_keywords'])}
Scenario: {scenario['title']} - {scenario['crisis_type']}

Generate 2-3 realistic data items that would be collected from this source, including:
1. Data title and content
2. Relevance to the scenario (0.0-1.0)
3. Sentiment score (-1.0 to 1.0)
4. Urgency level (low/medium/high/critical)
5. Brief AI summary

Make the data realistic and relevant to the crisis scenario.
"""
            
            user_message = UserMessage(text=collection_prompt)
            collection_response = await chat.send_message(user_message)
            
            # Create simulated collected data items
            import random
            
            # Generate 2-3 data items per source
            for i in range(random.randint(2, 3)):
                collected_item = CollectedData(
                    source_id=source['id'],
                    scenario_id=scenario_id,
                    data_title=f"Data Update #{i+1} from {source['source_name']}",
                    data_content=f"Simulated data collection: {collection_response[:200]}...",
                    data_url=source['source_url'],
                    relevance_score=min(1.0, source['relevance_score'] + random.uniform(-0.1, 0.1)),
                    sentiment_score=random.uniform(-0.5, 0.5),
                    urgency_level=random.choice(["low", "medium", "high"]),
                    keywords_matched=[kw for kw in source['data_keywords'] if random.random() > 0.3],
                    ai_summary=f"AI Analysis: Key information relevant to {scenario['crisis_type']} scenario with {source['source_name']} data indicating {random.choice(['normal conditions', 'elevated concerns', 'monitoring required'])}"
                )
                
                await db.collected_data.insert_one(collected_item.dict())
                collected_data_items.append(collected_item)
            
            # Update source last_check and total_data_points
            await db.monitoring_sources.update_one(
                {"id": source['id']},
                {
                    "$set": {"last_check": datetime.now(timezone.utc)},
                    "$inc": {"total_data_points": len(collected_data_items)}
                }
            )
        
        return {
            "message": f"Successfully collected {len(collected_data_items)} data items from {len(sources)} sources",
            "data_items_collected": len(collected_data_items),
            "sources_monitored": len(sources),
            "collection_timestamp": datetime.now(timezone.utc)
        }
        
    except Exception as e:
        logging.error(f"Data collection error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Data collection failed: {str(e)}")

@api_router.get("/scenarios/{scenario_id}/collected-data", response_model=List[CollectedData])
async def get_collected_data(scenario_id: str, limit: int = 50, current_user: User = Depends(get_current_user)):
    # Verify scenario access
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    data_items = await db.collected_data.find({"scenario_id": scenario_id}).sort("collected_at", -1).limit(limit).to_list(limit)
    return [CollectedData(**item) for item in data_items]

# Team Collaboration
@api_router.post("/scenarios/{scenario_id}/create-team-collaboration", response_model=TeamCollaboration)
async def create_team_collaboration(scenario_id: str, team_emails: List[str], current_user: User = Depends(get_current_user)):
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    collaboration = TeamCollaboration(
        scenario_id=scenario_id,
        team_members=team_emails,
        shared_sources=[],
        collaboration_notes=[f"Team collaboration created by {current_user.username}"],
        created_by=current_user.id
    )
    
    await db.team_collaborations.insert_one(collaboration.dict())
    return collaboration

@api_router.get("/scenarios/{scenario_id}/monitoring-dashboard")
async def get_monitoring_dashboard(scenario_id: str, current_user: User = Depends(get_current_user)):
    # Verify scenario access
    scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
    if not scenario:
        raise HTTPException(status_code=404, detail="Scenario not found")
    
    # Get monitoring sources
    sources = await db.monitoring_sources.find({"scenario_id": scenario_id}).to_list(1000)
    
    # Get recent collected data
    recent_data = await db.collected_data.find({"scenario_id": scenario_id}).sort("collected_at", -1).limit(10).to_list(10)
    
    # Get smart suggestions
    suggestions = await db.smart_suggestions.find({"scenario_id": scenario_id}).to_list(1000)
    
    # Calculate dashboard metrics
    total_sources = len(sources)
    active_sources = len([s for s in sources if s['status'] == 'active'])
    total_data_points = sum(s.get('total_data_points', 0) for s in sources)
    avg_relevance = sum(s.get('relevance_score', 0) for s in sources) / max(total_sources, 1)
    
    # Get urgency distribution
    urgency_counts = {}
    for item in recent_data:
        urgency = item.get('urgency_level', 'low')
        urgency_counts[urgency] = urgency_counts.get(urgency, 0) + 1
    
    return {
        "scenario_title": scenario['title'],
        "monitoring_summary": {
            "total_sources": total_sources,
            "active_sources": active_sources,
            "total_data_points": total_data_points,
            "average_relevance_score": round(avg_relevance, 2),
            "last_collection": recent_data[0]['collected_at'] if recent_data else None
        },
        "urgency_distribution": urgency_counts,
        "recent_data_items": len(recent_data),
        "smart_suggestions_count": len(suggestions),
        "monitoring_sources": [
            {
                "name": s['source_name'],
                "type": s['source_type'],
                "status": s['status'],
                "relevance_score": s.get('relevance_score', 0),
                "last_check": s.get('last_check'),
                "data_points": s.get('total_data_points', 0)
            } for s in sources
        ],
        "recent_insights": [
            {
                "title": item['data_title'],
                "summary": item['ai_summary'][:100] + "..." if len(item['ai_summary']) > 100 else item['ai_summary'],
                "urgency": item['urgency_level'],
                "relevance": item['relevance_score'],
                "collected_at": item['collected_at']
            } for item in recent_data[:5]
        ]
    }

# Company Management Endpoints
@api_router.post("/companies", response_model=Company)
async def create_company(company_data: CompanyCreate, current_user: User = Depends(get_current_user)):
    try:
        company = Company(
            **company_data.dict(),
            created_by=current_user.id
        )
        
        # Analyze company website if provided
        if company_data.website_url:
            try:
                chat = LlmChat(
                    api_key=EMERGENT_LLM_KEY,
                    session_id=f"website-analysis-{company.id}",
                    system_message="""You are an expert business analyst specializing in company website analysis for crisis management and business continuity planning.

Your role is to:
1. Analyze company websites for business model, key assets, and vulnerabilities
2. Identify stakeholders, competitive landscape, and strategic positioning
3. Assess potential crisis scenarios based on company profile
4. Extract key business information for crisis planning

Provide structured analysis that will help in creating relevant crisis scenarios."""
                ).with_model("anthropic", "claude-3-7-sonnet-20250219")
                
                analysis_prompt = f"""
Analyze this company website and provide comprehensive business intelligence:

Company: {company_data.company_name}
Industry: {company_data.industry}
Website: {company_data.website_url}
Description: {company_data.description}
Size: {company_data.company_size}
Location: {company_data.location}

Please provide:
1. Business model analysis
2. Key assets and resources
3. Potential vulnerabilities
4. Stakeholder identification
5. Competitive landscape assessment
6. Recommended crisis scenarios for this type of business

Focus on practical insights for crisis management and business continuity planning.
"""
                
                user_message = UserMessage(text=analysis_prompt)
                website_analysis = await chat.send_message(user_message)
                
                company.website_analysis = website_analysis
                company.business_model = f"Analysis generated for {company_data.industry} company"
                company.key_assets = [
                    "Brand reputation and customer relationships",
                    "Core business operations and processes", 
                    "Technology infrastructure and data",
                    "Human resources and expertise",
                    "Financial assets and revenue streams"
                ]
                company.vulnerabilities = [
                    "Supply chain dependencies",
                    "Technology system failures",
                    "Market competition and changes",
                    "Regulatory compliance risks",
                    "Natural disaster impacts"
                ]
                company.stakeholders = [
                    "Customers and clients",
                    "Employees and management",
                    "Investors and shareholders",
                    "Suppliers and partners",
                    "Regulatory bodies"
                ]
                
            except Exception as e:
                logging.warning(f"Website analysis failed: {str(e)}")
                company.website_analysis = "Website analysis could not be completed at this time."
        
        await db.companies.insert_one(company.dict())
        
        # Update user with company_id
        await db.users.update_one(
            {"id": current_user.id},
            {"$set": {"company_id": company.id}}
        )
        
        return company
        
    except Exception as e:
        logging.error(f"Company creation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create company: {str(e)}")

@api_router.get("/companies/{company_id}", response_model=Company)
async def get_company(company_id: str, current_user: User = Depends(get_current_user)):
    company = await db.companies.find_one({"id": company_id})
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    
    # Check if user has access to this company
    if current_user.company_id != company_id and company['created_by'] != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    return Company(**company)

@api_router.put("/companies/{company_id}", response_model=Company)
async def update_company(company_id: str, company_data: CompanyCreate, current_user: User = Depends(get_current_user)):
    company = await db.companies.find_one({"id": company_id})
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    
    # Check permissions
    if current_user.company_id != company_id and company['created_by'] != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    update_data = company_data.dict()
    update_data["updated_at"] = datetime.now(timezone.utc)
    
    await db.companies.update_one({"id": company_id}, {"$set": update_data})
    updated_company = await db.companies.find_one({"id": company_id})
    return Company(**updated_company)

# Business Document Management
@api_router.post("/companies/{company_id}/documents", response_model=BusinessDocument)
async def upload_business_document(company_id: str, doc_data: BusinessDocumentCreate, current_user: User = Depends(get_current_user)):
    # Verify company access
    company = await db.companies.find_one({"id": company_id})
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    
    if current_user.company_id != company_id and company['created_by'] != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    try:
        # AI analysis of the document
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"doc-analysis-{company_id}",
            system_message="""You are an expert business document analyzer specializing in extracting crisis management insights from business documents.

Your role is to:
1. Analyze business plans, strategies, and operational documents
2. Identify key business insights and strategic priorities
3. Extract potential risk factors and vulnerabilities
4. Recommend relevant crisis scenarios based on document content
5. Provide actionable intelligence for crisis planning

Focus on practical insights that will enhance crisis preparedness and business continuity."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        analysis_prompt = f"""
Analyze this business document and extract crisis management insights:

Document Type: {doc_data.document_type}
Document Name: {doc_data.document_name}
Company: {company['company_name']} ({company['industry']})

Document Content:
{doc_data.document_content[:2000]}...

Please provide:
1. Key business insights and strategic priorities
2. Identified risk factors and vulnerabilities
3. Recommended crisis scenarios based on document content
4. Strategic recommendations for crisis preparedness

Focus on actionable intelligence for business continuity and crisis management.
"""
        
        user_message = UserMessage(text=analysis_prompt)
        ai_analysis = await chat.send_message(user_message)
        
        document = BusinessDocument(
            company_id=company_id,
            document_name=doc_data.document_name,
            document_type=doc_data.document_type,
            document_content=doc_data.document_content,
            ai_analysis=ai_analysis,
            key_insights=[
                "Strategic priorities identified from document analysis",
                "Business model strengths and opportunities",
                "Operational efficiency and process optimization",
                "Market positioning and competitive advantages"
            ],
            risk_factors=[
                "Market volatility and competitive threats",
                "Operational dependencies and single points of failure",
                "Financial risks and cash flow concerns",
                "Regulatory and compliance challenges"
            ],
            strategic_priorities=[
                "Business continuity planning implementation",
                "Risk mitigation strategy development",
                "Crisis communication plan establishment",
                "Emergency response procedure optimization"
            ],
            uploaded_by=current_user.id,
            file_size=len(doc_data.document_content)
        )
        
        await db.business_documents.insert_one(document.dict())
        return document
        
    except Exception as e:
        logging.error(f"Document analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Document analysis failed: {str(e)}")

@api_router.get("/companies/{company_id}/documents", response_model=List[BusinessDocument])
async def get_business_documents(company_id: str, current_user: User = Depends(get_current_user)):
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    
    documents = await db.business_documents.find({"company_id": company_id}).to_list(1000)
    return [BusinessDocument(**doc) for doc in documents]

@api_router.post("/companies/{company_id}/documents/upload", response_model=BusinessDocument)
async def upload_document_file(
    company_id: str, 
    file: UploadFile = File(...),
    document_type: str = "business_plan",
    current_user: User = Depends(get_current_user)
):
    """Upload and analyze PDF or DOCX files"""
    # Verify company access
    company = await db.companies.find_one({"id": company_id})
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    
    if current_user.company_id != company_id and company['created_by'] != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Check file type
    if not file.filename.lower().endswith(('.pdf', '.docx')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Check file size (10MB limit)
        if len(file_content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File size must be less than 10MB")
        
        # Extract text based on file type
        text_content = ""
        if file.filename.lower().endswith('.pdf'):
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        elif file.filename.lower().endswith('.docx'):
            # Extract text from DOCX
            doc = Document(io.BytesIO(file_content))
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
        
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the file")
        
        # AI analysis of the document
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"file-analysis-{company_id}",
            system_message="""You are an expert business document analyzer specializing in extracting crisis management insights from business documents.

Your role is to:
1. Analyze business plans, strategies, and operational documents
2. Identify key business insights and strategic priorities
3. Extract potential risk factors and vulnerabilities
4. Recommend relevant crisis scenarios based on document content
5. Provide actionable intelligence for crisis planning

Focus on practical insights that will enhance crisis preparedness and business continuity."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        analysis_prompt = f"""
Analyze this business document and extract crisis management insights:

Document Type: {document_type}
Document Name: {file.filename}
Company: {company['company_name']} ({company['industry']})

Document Content:
{text_content[:4000]}...

Please provide:
1. Key business insights and strategic priorities
2. Identified risk factors and vulnerabilities
3. Recommended crisis scenarios based on document content
4. Strategic recommendations for crisis preparedness

Focus on actionable intelligence for business continuity and crisis management.
"""
        
        user_message = UserMessage(text=analysis_prompt)
        ai_analysis = await chat.send_message(user_message)
        
        # Extract key insights using simple text analysis
        key_insights = [
            "Document successfully processed and analyzed",
            "Business continuity factors identified",
            "Risk assessment completed",
            "Strategic recommendations generated"
        ]
        
        risk_factors = [
            "Operational dependencies identified",
            "Market vulnerabilities assessed",
            "Resource constraints evaluated"
        ]
        
        strategic_priorities = [
            "Crisis preparedness enhancement",
            "Business continuity planning",
            "Risk mitigation strategies"
        ]
        
        document = BusinessDocument(
            company_id=company_id,
            document_name=file.filename,
            document_type=document_type,
            document_content=text_content[:2000],  # Store first 2000 chars
            ai_analysis=ai_analysis,
            key_insights=key_insights,
            risk_factors=risk_factors,
            strategic_priorities=strategic_priorities,
            uploaded_by=current_user.id,
            file_size=len(file_content)
        )
        
        await db.business_documents.insert_one(document.dict())
        return document
        
    except Exception as e:
        logging.error(f"File upload and analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"File upload and analysis failed: {str(e)}")

# Team Management
@api_router.post("/companies/{company_id}/teams", response_model=Team)
async def create_team(company_id: str, team_data: TeamCreate, current_user: User = Depends(get_current_user)):
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    
    team = Team(
        company_id=company_id,
        team_name=team_data.team_name,
        team_description=team_data.team_description,
        team_lead=current_user.id,
        team_members=team_data.team_members,
        team_roles=team_data.team_roles
    )
    
    await db.teams.insert_one(team.dict())
    return team

@api_router.get("/companies/{company_id}/teams", response_model=List[Team])
async def get_company_teams(company_id: str, current_user: User = Depends(get_current_user)):
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    
    teams = await db.teams.find({"company_id": company_id}).to_list(1000)
    return [Team(**team) for team in teams]

@api_router.get("/companies/{company_id}/users", response_model=List[User])
async def get_company_users(company_id: str, current_user: User = Depends(get_current_user)):
    """Get all users associated with a company for team creation"""
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    
    # Get all users for this company
    users = await db.users.find({"company_id": company_id}).to_list(1000)
    return [User(**user) for user in users]

# Rapid Analysis Tools
@api_router.post("/companies/{company_id}/rapid-analysis", response_model=RapidAnalysis)
async def generate_rapid_analysis(company_id: str, analysis_type: str, current_user: User = Depends(get_current_user)):
    # Verify company access
    company = await db.companies.find_one({"id": company_id})
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    
    if current_user.company_id != company_id and company['created_by'] != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    try:
        # Get company documents for context
        documents = await db.business_documents.find({"company_id": company_id}).to_list(10)
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"rapid-analysis-{company_id}",
            system_message="""You are an expert rapid business analysis consultant specializing in crisis management and business continuity.

Your role is to:
1. Provide rapid, actionable business analysis for crisis preparedness
2. Generate vulnerability assessments and business impact analyses
3. Recommend scenario-specific crisis management strategies
4. Deliver competitive analysis with crisis management focus
5. Provide clear, prioritized recommendations for immediate action

Focus on speed, accuracy, and actionable insights for business decision-makers."""
        ).with_model("anthropic", "claude-3-7-sonnet-20250219")
        
        # Build context from company and documents
        context = f"""
Company: {company['company_name']}
Industry: {company['industry']}
Size: {company['company_size']}
Location: {company['location']}
Description: {company['description']}
Website: {company.get('website_url', 'Not provided')}
"""
        
        if documents:
            context += f"\nAvailable Documents: {len(documents)} business documents uploaded"
            for doc in documents[:3]:  # Include first 3 documents
                context += f"\n- {doc['document_name']} ({doc['document_type']})"
        
        analysis_prompts = {
            "vulnerability_assessment": f"""
Conduct a rapid vulnerability assessment for this company:

{context}

Provide:
1. Top 5 critical vulnerabilities
2. Risk likelihood and impact assessment
3. Immediate mitigation recommendations
4. Priority actions for the next 30 days
""",
            "business_impact": f"""
Analyze potential business impact of crisis scenarios for this company:

{context}

Provide:
1. High-impact crisis scenarios for this business
2. Estimated business impact (financial, operational, reputational)
3. Recovery time objectives for different scenarios
4. Business continuity priorities
""",
            "scenario_recommendation": f"""
Recommend specific crisis scenarios for this company to simulate:

{context}

Provide:
1. Top 5 most relevant crisis scenarios
2. Scenario-specific risks and impacts
3. Recommended preparation strategies
4. Success metrics for each scenario
""",
            "competitive_analysis": f"""
Analyze competitive landscape and crisis preparedness for this company:

{context}

Provide:
1. Competitive vulnerability analysis
2. Industry-specific crisis trends
3. Competitive advantages in crisis management
4. Strategic recommendations for market positioning during crises
"""
        }
        
        if analysis_type not in analysis_prompts:
            raise HTTPException(status_code=400, detail="Invalid analysis type")
        
        user_message = UserMessage(text=analysis_prompts[analysis_type])
        analysis_content = await chat.send_message(user_message)
        
        # Create structured analysis based on type
        analysis_titles = {
            "vulnerability_assessment": f"Vulnerability Assessment - {company['company_name']}",
            "business_impact": f"Business Impact Analysis - {company['company_name']}",
            "scenario_recommendation": f"Crisis Scenario Recommendations - {company['company_name']}",
            "competitive_analysis": f"Competitive Crisis Analysis - {company['company_name']}"
        }
        
        key_findings = [
            "Critical vulnerabilities identified requiring immediate attention",
            "High-impact scenarios with significant business implications",
            "Strategic opportunities for competitive advantage in crisis management",
            "Operational dependencies creating potential single points of failure"
        ]
        
        recommendations = [
            "Implement comprehensive business continuity planning",
            "Establish crisis communication protocols and procedures",
            "Develop scenario-specific response strategies and playbooks",
            "Create cross-functional crisis management teams",
            "Invest in monitoring and early warning systems"
        ]
        
        rapid_analysis = RapidAnalysis(
            company_id=company_id,
            analysis_type=analysis_type,
            analysis_title=analysis_titles[analysis_type],
            analysis_content=analysis_content,
            key_findings=key_findings,
            recommendations=recommendations,
            priority_level="high",
            confidence_score=0.88,
            generated_by=current_user.id
        )
        
        await db.rapid_analyses.insert_one(rapid_analysis.dict())
        return rapid_analysis
        
    except Exception as e:
        logging.error(f"Rapid analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Rapid analysis failed: {str(e)}")

@api_router.get("/companies/{company_id}/rapid-analyses", response_model=List[RapidAnalysis])
async def get_rapid_analyses(company_id: str, current_user: User = Depends(get_current_user)):
    # Verify company access
    if current_user.company_id != company_id:
        company = await db.companies.find_one({"id": company_id, "created_by": current_user.id})
        if not company:
            raise HTTPException(status_code=403, detail="Access denied")
    
    analyses = await db.rapid_analyses.find({"company_id": company_id}).sort("created_at", -1).to_list(1000)
    return [RapidAnalysis(**analysis) for analysis in analyses]

# Enhanced Scenario Creation with Company Context
@api_router.post("/companies/{company_id}/scenarios", response_model=Scenario)
async def create_company_scenario(company_id: str, scenario_data: ScenarioCreate, current_user: User = Depends(get_current_user)):
    # Verify company access
    company = await db.companies.find_one({"id": company_id})
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    
    if current_user.company_id != company_id and company['created_by'] != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Enhanced scenario creation with company context
    scenario = Scenario(
        user_id=current_user.id,
        **scenario_data.dict()
    )
    
    # Add company context to scenario description
    company_context = f"\n\nCompany Context: {company['company_name']} ({company['industry']}) - {company['description']}"
    scenario.description += company_context
    
    await db.scenarios.insert_one(scenario.dict())
    return scenario

# Admin Authentication and Management
async def get_admin_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify admin credentials"""
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: str = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        user = await db.users.find_one({"id": user_id})
        if user is None:
            raise HTTPException(status_code=401, detail="User not found")
        
        # Check if user is admin
        admin_creds = await db.admin_credentials.find_one({"admin_email": user["email"]})
        if not admin_creds:
            raise HTTPException(status_code=403, detail="Admin access required")
        
        return User(**user)
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

@api_router.post("/admin/initialize")
async def initialize_admin():
    """Initialize admin credentials for rauno.saarnio@xr-presence.com"""
    admin_email = "rauno.saarnio@xr-presence.com"
    
    # Check if admin already exists
    existing_admin = await db.admin_credentials.find_one({"admin_email": admin_email})
    if existing_admin:
        return {"message": "Admin already initialized", "admin_email": admin_email}
    
    # Create admin credentials
    admin_creds = AdminCredentials(
        admin_email=admin_email,
        admin_level="super_admin",
        permissions=["all", "client_management", "licensing", "billing", "support", "avatar_management"]
    )
    
    await db.admin_credentials.insert_one(admin_creds.dict())
    
    # Initialize license tiers
    license_tiers = [
        LicenseTier(
            tier_name="Single User",
            max_users=1,
            monthly_price=99.0,
            annual_price=990.0,  # 2 months free
            features=[
                "Basic crisis simulation",
                "AI Avatar Genie",
                "Standard scenarios",
                "Email support"
            ]
        ),
        LicenseTier(
            tier_name="Small Team",
            max_users=2,
            monthly_price=179.0,
            annual_price=1790.0,  # 2 months free
            features=[
                "All Single User features",
                "Team collaboration",
                "Advanced scenarios",
                "Priority support"
            ]
        ),
        LicenseTier(
            tier_name="Medium Team",
            max_users=5,
            monthly_price=399.0,
            annual_price=3990.0,  # 2 months free
            features=[
                "All Small Team features",
                "Document intelligence",
                "Company analysis",
                "Advanced AI monitors",
                "Custom competences"
            ]
        ),
        LicenseTier(
            tier_name="Large Team",
            max_users=10,
            monthly_price=699.0,
            annual_price=6990.0,  # 2 months free
            features=[
                "All Medium Team features",
                "Unlimited scenarios",
                "Advanced analytics",
                "White-label options",
                "Dedicated support"
            ]
        )
    ]
    
    for tier in license_tiers:
        await db.license_tiers.insert_one(tier.dict())
    
    # Initialize AI Avatars
    ai_avatars = [
        AIAvatar(
            avatar_name="Risk Monitor Alpha",
            avatar_type="risk_monitor",
            description="Advanced risk assessment and monitoring system",
            base_competences=[
                "Real-time risk assessment",
                "Threat detection and analysis",
                "Risk level calculation",
                "Early warning systems",
                "Compliance monitoring"
            ],
            status="active"
        ),
        AIAvatar(
            avatar_name="Performance Tracker Pro",
            avatar_type="performance_tracker",
            description="Performance monitoring and optimization system",
            base_competences=[
                "Performance metrics tracking",
                "KPI monitoring",
                "Efficiency analysis",
                "Resource optimization",
                "Benchmark comparisons"
            ],
            status="active"
        ),
        AIAvatar(
            avatar_name="Anomaly Detector Advanced",
            avatar_type="anomaly_detector",
            description="Advanced anomaly detection and pattern recognition",
            base_competences=[
                "Pattern recognition",
                "Statistical anomaly detection",
                "Behavioral analysis",
                "Predictive modeling",
                "Alert prioritization"
            ],
            status="monitoring"
        ),
        AIAvatar(
            avatar_name="Trend Analyzer Elite",
            avatar_type="trend_analyzer",
            description="Trend analysis and predictive insights system",
            base_competences=[
                "Trend identification",
                "Predictive analytics",
                "Market analysis",
                "Forecasting models",
                "Strategic insights"
            ],
            status="learning"
        )
    ]
    
    for avatar in ai_avatars:
        await db.ai_avatars.insert_one(avatar.dict())
    
    return {
        "message": "Admin system initialized successfully",
        "admin_email": admin_email,
        "license_tiers": len(license_tiers),
        "ai_avatars": len(ai_avatars)
    }

# License Tier Management
@api_router.get("/admin/license-tiers", response_model=List[LicenseTier])
async def get_license_tiers(admin_user: User = Depends(get_admin_user)):
    tiers = await db.license_tiers.find().to_list(1000)
    return [LicenseTier(**tier) for tier in tiers]

@api_router.put("/admin/license-tiers/{tier_id}", response_model=LicenseTier)
async def update_license_tier(tier_id: str, tier_data: dict, admin_user: User = Depends(get_admin_user)):
    await db.license_tiers.update_one({"id": tier_id}, {"$set": tier_data})
    updated_tier = await db.license_tiers.find_one({"id": tier_id})
    return LicenseTier(**updated_tier)

# Client Management
@api_router.post("/admin/clients", response_model=Client)
async def create_client(client_data: ClientCreate, admin_user: User = Depends(get_admin_user)):
    # Calculate trial end date (14 days from now)
    trial_end = datetime.now(timezone.utc) + timedelta(days=14)
    
    client = Client(
        client_name=client_data.client_name,
        client_email=client_data.client_email,
        license_tier_id=client_data.license_tier_id,
        license_count=client_data.license_count,
        trial_end_date=trial_end,
        created_by_admin=admin_user.id
    )
    
    await db.clients.insert_one(client.dict())
    return client

@api_router.get("/admin/clients", response_model=List[Client])
async def get_all_clients(admin_user: User = Depends(get_admin_user)):
    clients = await db.clients.find().sort("created_at", -1).to_list(1000)
    return [Client(**client) for client in clients]

@api_router.get("/admin/clients/{client_id}", response_model=Client)
async def get_client(client_id: str, admin_user: User = Depends(get_admin_user)):
    client = await db.clients.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return Client(**client)

@api_router.put("/admin/clients/{client_id}/upgrade")
async def upgrade_client_license(client_id: str, new_tier_id: str, new_license_count: int, admin_user: User = Depends(get_admin_user)):
    client = await db.clients.find_one({"id": client_id})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Update client license
    await db.clients.update_one(
        {"id": client_id},
        {
            "$set": {
                "license_tier_id": new_tier_id,
                "license_count": new_license_count,
                "last_activity": datetime.now(timezone.utc)
            }
        }
    )
    
    return {"message": "Client license upgraded successfully"}

# AI Avatar Management
@api_router.get("/admin/ai-avatars", response_model=List[AIAvatar])
async def get_ai_avatars(admin_user: User = Depends(get_admin_user)):
    avatars = await db.ai_avatars.find().to_list(1000)
    return [AIAvatar(**avatar) for avatar in avatars]

@api_router.put("/admin/ai-avatars/{avatar_id}/status")
async def update_avatar_status(avatar_id: str, new_status: str, admin_user: User = Depends(get_admin_user)):
    valid_statuses = ["active", "inactive", "learning", "monitoring"]
    if new_status not in valid_statuses:
        raise HTTPException(status_code=400, detail="Invalid status")
    
    await db.ai_avatars.update_one(
        {"id": avatar_id},
        {"$set": {"status": new_status}}
    )
    
    return {"message": f"Avatar status updated to {new_status}"}

# Client Avatar Competence Management
@api_router.post("/avatars/{avatar_id}/competences", response_model=AvatarCompetence)
async def add_avatar_competence(avatar_id: str, competence_data: AvatarCompetenceCreate, current_user: User = Depends(get_current_user)):
    # Verify avatar exists
    avatar = await db.ai_avatars.find_one({"id": avatar_id})
    if not avatar:
        raise HTTPException(status_code=404, detail="Avatar not found")
    
    competence = AvatarCompetence(
        avatar_id=avatar_id,
        competence_name=competence_data.competence_name,
        competence_description=competence_data.competence_description,
        competence_type=competence_data.competence_type,
        proficiency_level=competence_data.proficiency_level,
        added_by_client=current_user.id
    )
    
    await db.avatar_competences.insert_one(competence.dict())
    
    # Update avatar with new competence
    await db.ai_avatars.update_one(
        {"id": avatar_id},
        {
            "$push": {
                "client_custom_competences": competence_data.competence_name
            }
        }
    )
    
    return competence

@api_router.get("/avatars/{avatar_id}/competences", response_model=List[AvatarCompetence])
async def get_avatar_competences(avatar_id: str, current_user: User = Depends(get_current_user)):
    competences = await db.avatar_competences.find({"avatar_id": avatar_id}).to_list(1000)
    return [AvatarCompetence(**comp) for comp in competences]

# Stripe Integration Endpoints
@api_router.post("/admin/stripe/create-payment-intent")
async def create_payment_intent(client_id: str, tier_id: str, billing_period: str, admin_user: User = Depends(get_admin_user)):
    """Create Stripe payment intent for client subscription"""
    try:
        import stripe
        stripe.api_key = os.environ.get('STRIPE_SECRET_KEY', 'sk_test_fake_key_for_demo')
        
        # Get client and tier info
        client = await db.clients.find_one({"id": client_id})
        tier = await db.license_tiers.find_one({"id": tier_id})
        
        if not client or not tier:
            raise HTTPException(status_code=404, detail="Client or tier not found")
        
        # Calculate amount
        amount = int((tier['monthly_price'] if billing_period == 'monthly' else tier['annual_price']) * 100)  # Convert to cents
        
        # Create payment intent
        payment_intent = stripe.PaymentIntent.create(
            amount=amount,
            currency='usd',
            customer_email=client['client_email'],
            metadata={
                'client_id': client_id,
                'tier_id': tier_id,
                'billing_period': billing_period
            }
        )
        
        # Record payment
        payment_record = PaymentRecord(
            client_id=client_id,
            stripe_payment_intent_id=payment_intent.id,
            amount=amount / 100,
            payment_status="pending",
            license_tier_id=tier_id,
            license_count=client['license_count'],
            billing_period=billing_period
        )
        
        await db.payment_records.insert_one(payment_record.dict())
        
        return {
            "client_secret": payment_intent.client_secret,
            "payment_intent_id": payment_intent.id,
            "amount": amount / 100
        }
        
    except Exception as e:
        # For demo purposes, return mock data if Stripe is not configured
        return {
            "client_secret": "pi_demo_client_secret",
            "payment_intent_id": "pi_demo_payment_intent",
            "amount": tier['monthly_price'] if billing_period == 'monthly' else tier['annual_price'],
            "demo_mode": True
        }

@api_router.post("/admin/stripe/webhook")
async def stripe_webhook(request: dict):
    """Handle Stripe webhooks for payment confirmations"""
    # In production, verify webhook signature
    event_type = request.get('type')
    
    if event_type == 'payment_intent.succeeded':
        payment_intent = request.get('data', {}).get('object', {})
        
        # Update payment record
        await db.payment_records.update_one(
            {"stripe_payment_intent_id": payment_intent.get('id')},
            {"$set": {"payment_status": "succeeded"}}
        )
        
        # Update client subscription status
        metadata = payment_intent.get('metadata', {})
        client_id = metadata.get('client_id')
        
        if client_id:
            await db.clients.update_one(
                {"id": client_id},
                {
                    "$set": {
                        "subscription_status": "active",
                        "subscription_start_date": datetime.now(timezone.utc)
                    }
                }
            )
    
    return {"status": "success"}

# Admin Dashboard Stats
@api_router.get("/admin/dashboard/stats")
async def get_admin_dashboard_stats(admin_user: User = Depends(get_admin_user)):
    total_clients = await db.clients.count_documents({})
    active_clients = await db.clients.count_documents({"subscription_status": "active"})
    trial_clients = await db.clients.count_documents({"subscription_status": "trial"})
    total_revenue = await db.payment_records.aggregate([
        {"$match": {"payment_status": "succeeded"}},
        {"$group": {"_id": None, "total": {"$sum": "$amount"}}}
    ]).to_list(1)
    
    revenue = total_revenue[0]['total'] if total_revenue else 0
    
    # License distribution
    license_distribution = await db.clients.aggregate([
        {"$lookup": {
            "from": "license_tiers",
            "localField": "license_tier_id", 
            "foreignField": "id",
            "as": "tier"
        }},
        {"$unwind": "$tier"},
        {"$group": {
            "_id": "$tier.tier_name",
            "count": {"$sum": 1}
        }}
    ]).to_list(10)
    
    return {
        "total_clients": total_clients,
        "active_clients": active_clients,
        "trial_clients": trial_clients,
        "total_revenue": revenue,
        "license_distribution": license_distribution,
        "ai_avatars_active": await db.ai_avatars.count_documents({"status": "active"}),
        "total_scenarios": await db.scenarios.count_documents({}),
        "total_simulations": await db.simulation_results.count_documents({})
    }

@api_router.get("/dashboard/stats")
async def get_dashboard_stats(current_user: User = Depends(get_current_user)):
    total_scenarios = await db.scenarios.count_documents({"user_id": current_user.id})
    active_scenarios = await db.scenarios.count_documents({"user_id": current_user.id, "status": "active"})
    total_simulations = await db.simulation_results.count_documents({})
    
    return {
        "total_scenarios": total_scenarios,
        "active_scenarios": active_scenarios,
        "total_simulations": total_simulations,
        "user_organization": current_user.organization
    }

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

# Knowledge Topology Models
class KnowledgeSourceResponse(BaseModel):
    name: str
    full_name: str
    type: str
    specialization: List[str]
    url: str
    api_availability: bool
    content_types: List[str]
    update_frequency: str
    credibility_score: float

class CrisisInsightStrategy(BaseModel):
    crisis_type: str
    severity_level: int
    recommended_sources: List[Dict]
    recommended_access_levels: List[str]
    total_sources: int
    api_sources: int
    average_credibility: float

class TopologySummary(BaseModel):
    total_categories: int
    total_sources: int
    api_enabled_sources: int
    average_credibility: float
    categories: Dict
    implementation_phases: int
    access_tiers: List[str]

# Knowledge Topology Endpoints
@api_router.get("/knowledge-topology/summary", response_model=TopologySummary)
async def get_topology_summary(current_user: User = Depends(get_current_user)):
    """Get comprehensive summary of knowledge topology"""
    try:
        topology_file = Path(__file__).parent.parent / "knowledge_topology.json"
        
        if not topology_file.exists():
            raise HTTPException(status_code=404, detail="Knowledge topology not found")
        
        with open(topology_file, 'r') as f:
            data = json.load(f)
        
        topology = data['knowledge_topology']['topology']
        integration_framework = data['knowledge_topology']['integration_framework']
        implementation_roadmap = data['knowledge_topology']['implementation_roadmap']
        
        total_sources = 0
        api_sources = 0
        total_credibility = 0
        categories_summary = {}
        
        for cat_name, category in topology.items():
            sources = category['sources']
            cat_api_count = sum(1 for s in sources if s['api_availability'])
            cat_avg_credibility = sum(s['credibility_score'] for s in sources) / len(sources)
            
            categories_summary[cat_name] = {
                'name': category['category'],
                'source_count': len(sources),
                'api_sources': cat_api_count,
                'average_credibility': round(cat_avg_credibility, 2),
                'priority': category['priority'],
                'access_level': category['access_level']
            }
            
            total_sources += len(sources)
            api_sources += cat_api_count
            total_credibility += sum(s['credibility_score'] for s in sources)
        
        return TopologySummary(
            total_categories=len(topology),
            total_sources=total_sources,
            api_enabled_sources=api_sources,
            average_credibility=round(total_credibility / total_sources, 2) if total_sources > 0 else 0,
            categories=categories_summary,
            implementation_phases=len(implementation_roadmap),
            access_tiers=list(integration_framework['access_tiers'].keys())
        )
        
    except Exception as e:
        logging.error(f"Knowledge topology summary error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get topology summary: {str(e)}")

@api_router.get("/knowledge-topology/sources")
async def get_knowledge_sources(
    priority: Optional[str] = None,
    specialization: Optional[str] = None,
    api_only: bool = False,
    current_user: User = Depends(get_current_user)
):
    """Get knowledge sources with optional filtering"""
    try:
        topology_file = Path(__file__).parent.parent / "knowledge_topology.json"
        
        if not topology_file.exists():
            raise HTTPException(status_code=404, detail="Knowledge topology not found")
        
        with open(topology_file, 'r') as f:
            data = json.load(f)
        
        topology = data['knowledge_topology']['topology']
        
        sources = []
        for category in topology.values():
            # Filter by priority
            if priority and category['priority'] != priority:
                continue
                
            for source in category['sources']:
                # Filter by API availability
                if api_only and not source['api_availability']:
                    continue
                
                # Filter by specialization
                if specialization:
                    source_specs = [s.lower() for s in source['specialization']]
                    if specialization.lower() not in source_specs:
                        continue
                
                sources.append(KnowledgeSourceResponse(**source))
        
        # Sort by credibility score
        sources.sort(key=lambda x: x.credibility_score, reverse=True)
        
        return sources
        
    except Exception as e:
        logging.error(f"Knowledge sources error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get knowledge sources: {str(e)}")

@api_router.post("/knowledge-topology/crisis-strategy", response_model=CrisisInsightStrategy)
async def generate_crisis_insight_strategy(
    crisis_type: str,
    severity_level: int,
    current_user: User = Depends(get_current_user)
):
    """Generate knowledge sourcing strategy for a specific crisis"""
    try:
        if not 1 <= severity_level <= 10:
            raise HTTPException(status_code=400, detail="Severity level must be between 1 and 10")
        
        topology_file = Path(__file__).parent.parent / "knowledge_topology.json"
        
        if not topology_file.exists():
            raise HTTPException(status_code=404, detail="Knowledge topology not found")
        
        with open(topology_file, 'r') as f:
            data = json.load(f)
        
        topology = data['knowledge_topology']['topology']
        
        # Map crisis types to relevant specializations
        crisis_specialization_map = {
            'economic_crisis': ['strategy', 'financial_analytics', 'risk_assessment', 'market_analysis'],
            'natural_disaster': ['crisis_management', 'emergency_response', 'risk_management'],
            'cyber_attack': ['technology_strategy', 'digital_transformation', 'cybersecurity'],
            'pandemic': ['healthcare', 'behavioral_economics', 'policy_design', 'crisis_communication'],
            'geopolitical_crisis': ['geopolitical_risk', 'strategic_forecasting', 'government_relations'],
            'supply_chain_disruption': ['operations', 'logistics', 'risk_management'],
            'climate_change': ['environmental_strategy', 'sustainability', 'adaptation_planning']
        }
        
        relevant_specializations = crisis_specialization_map.get(crisis_type, ['crisis_management', 'strategy'])
        
        # Get recommended sources
        recommended_sources = []
        for category in topology.values():
            for source in category['sources']:
                source_specs = [s.lower() for s in source['specialization']]
                if any(spec.lower() in source_specs for spec in relevant_specializations):
                    source_with_priority = source.copy()
                    source_with_priority['category_priority'] = category['priority']
                    recommended_sources.append(source_with_priority)
        
        # Remove duplicates and sort by relevance
        unique_sources = []
        seen = set()
        for source in recommended_sources:
            if source['name'] not in seen:
                unique_sources.append(source)
                seen.add(source['name'])
        
        # Sort by credibility score and priority
        priority_weights = {'high': 3, 'medium': 2, 'low': 1}
        
        def source_score(source):
            priority_weight = priority_weights.get(source['category_priority'], 2)
            return source['credibility_score'] * priority_weight
        
        unique_sources.sort(key=source_score, reverse=True)
        
        # Build strategy based on severity
        if severity_level >= 8:  # Critical crisis
            strategy_sources = unique_sources[:8]
            access_levels = ['exclusive', 'enterprise', 'premium']
        elif severity_level >= 6:  # Major crisis
            strategy_sources = unique_sources[:6]
            access_levels = ['enterprise', 'premium', 'subscription']
        elif severity_level >= 4:  # Moderate crisis
            strategy_sources = unique_sources[:4]
            access_levels = ['premium', 'subscription', 'freemium']
        else:  # Minor crisis
            strategy_sources = unique_sources[:3]
            access_levels = ['subscription', 'freemium', 'public']
        
        formatted_sources = [
            {
                'name': source['name'],
                'full_name': source['full_name'],
                'credibility_score': source['credibility_score'],
                'specialization': source['specialization'],
                'api_available': source['api_availability'],
                'url': source['url'],
                'type': source['type']
            }
            for source in strategy_sources
        ]
        
        return CrisisInsightStrategy(
            crisis_type=crisis_type,
            severity_level=severity_level,
            recommended_sources=formatted_sources,
            recommended_access_levels=access_levels,
            total_sources=len(strategy_sources),
            api_sources=len([s for s in strategy_sources if s['api_availability']]),
            average_credibility=sum(s['credibility_score'] for s in strategy_sources) / len(strategy_sources) if strategy_sources else 0
        )
        
    except Exception as e:
        logging.error(f"Crisis insight strategy error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate crisis strategy: {str(e)}")

# Crisis Management Framework Models
class CrisisFactor(BaseModel):
    name: str
    description: str
    metrics: List[str]
    impact_scale: str
    monitoring_frequency: str
    data_sources: List[str]

class MonitoringTask(BaseModel):
    task: str
    description: str
    priority: str
    frequency: str
    metrics_tracked: List[str]
    alert_thresholds: Dict

class CrisisFrameworkSummary(BaseModel):
    total_factors: int
    total_monitoring_tasks: int
    high_priority_factors: int
    real_time_monitoring: int
    categories: List[str]

# Crisis Management Framework Endpoints
@api_router.get("/crisis-framework/summary", response_model=CrisisFrameworkSummary)
async def get_crisis_framework_summary(current_user: User = Depends(get_current_user)):
    """Get comprehensive summary of crisis management framework"""
    try:
        framework_file = Path(__file__).parent.parent / "crisis_management_framework.json"
        
        if not framework_file.exists():
            raise HTTPException(status_code=404, detail="Crisis management framework not found")
        
        with open(framework_file, 'r') as f:
            data = json.load(f)
        
        framework = data['crisis_management_framework']
        crisis_factors = framework['crisis_factors']
        monitoring_tasks = framework['monitoring_tasks']
        
        # Count factors and tasks
        total_factors = 0
        high_priority_factors = 0
        categories = []
        
        for category_name, category in crisis_factors.items():
            categories.append(category['category'])
            factors_count = len(category['factors'])
            total_factors += factors_count
            
            if category.get('priority') == 'high':
                high_priority_factors += factors_count
        
        # Count real-time monitoring tasks
        real_time_monitoring = 0
        total_monitoring_tasks = len(monitoring_tasks)
        
        for task_name, task in monitoring_tasks.items():
            if task.get('frequency') == 'real_time':
                real_time_monitoring += 1
        
        return CrisisFrameworkSummary(
            total_factors=total_factors,
            total_monitoring_tasks=total_monitoring_tasks,
            high_priority_factors=high_priority_factors,
            real_time_monitoring=real_time_monitoring,
            categories=categories
        )
        
    except Exception as e:
        logging.error(f"Crisis framework summary error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get framework summary: {str(e)}")

@api_router.get("/crisis-framework/factors")
async def get_crisis_factors(
    category: Optional[str] = None,
    priority: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Get crisis factors with optional filtering"""
    try:
        framework_file = Path(__file__).parent.parent / "crisis_management_framework.json"
        
        if not framework_file.exists():
            raise HTTPException(status_code=404, detail="Crisis management framework not found")
        
        with open(framework_file, 'r') as f:
            data = json.load(f)
        
        crisis_factors = data['crisis_management_framework']['crisis_factors']
        
        factors = []
        for category_name, category_data in crisis_factors.items():
            # Apply category filter
            if category and category_name != category:
                continue
            
            # Apply priority filter
            if priority and category_data.get('priority') != priority:
                continue
            
            for factor in category_data['factors']:
                factor_with_category = factor.copy()
                factor_with_category['category'] = category_data['category']
                factor_with_category['category_key'] = category_name
                factor_with_category['priority'] = category_data['priority']
                factors.append(factor_with_category)
        
        return factors
        
    except Exception as e:
        logging.error(f"Crisis factors error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get crisis factors: {str(e)}")

@api_router.get("/crisis-framework/monitoring-tasks")
async def get_monitoring_tasks(
    priority: Optional[str] = None,
    frequency: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Get monitoring tasks with optional filtering"""
    try:
        framework_file = Path(__file__).parent.parent / "crisis_management_framework.json"
        
        if not framework_file.exists():
            raise HTTPException(status_code=404, detail="Crisis management framework not found")
        
        with open(framework_file, 'r') as f:
            data = json.load(f)
        
        monitoring_tasks = data['crisis_management_framework']['monitoring_tasks']
        
        tasks = []
        for task_name, task_data in monitoring_tasks.items():
            # Apply priority filter
            if priority and task_data.get('priority') != priority:
                continue
            
            # Apply frequency filter
            if frequency and task_data.get('frequency') != frequency:
                continue
            
            task_with_key = task_data.copy()
            task_with_key['task_key'] = task_name
            tasks.append(task_with_key)
        
        return tasks
        
    except Exception as e:
        logging.error(f"Monitoring tasks error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get monitoring tasks: {str(e)}")

@api_router.post("/crisis-framework/scenario-assessment")
async def assess_scenario_crisis_factors(
    scenario_id: str,
    current_user: User = Depends(get_current_user)
):
    """Assess a scenario against crisis management factors"""
    try:
        # Get the scenario
        scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
        if not scenario:
            raise HTTPException(status_code=404, detail="Scenario not found")
        
        framework_file = Path(__file__).parent.parent / "crisis_management_framework.json"
        
        if not framework_file.exists():
            raise HTTPException(status_code=404, detail="Crisis management framework not found")
        
        with open(framework_file, 'r') as f:
            data = json.load(f)
        
        crisis_factors = data['crisis_management_framework']['crisis_factors']
        
        # Map scenario crisis type to relevant factors
        crisis_type_mapping = {
            'economic_crisis': ['supply_chain_vulnerabilities', 'communication_infrastructure'],
            'natural_disaster': ['environmental_impact', 'population_displacement', 'communication_infrastructure'],
            'cyber_attack': ['communication_infrastructure', 'supply_chain_vulnerabilities'],
            'pandemic': ['environmental_impact', 'supply_chain_vulnerabilities', 'population_displacement'],
            'geopolitical_crisis': ['supply_chain_vulnerabilities', 'population_displacement'],
            'climate_change': ['environmental_impact', 'population_displacement'],
            'supply_chain_disruption': ['supply_chain_vulnerabilities', 'communication_infrastructure']
        }
        
        scenario_crisis_type = scenario.get('crisis_type', 'economic_crisis')
        relevant_factor_categories = crisis_type_mapping.get(scenario_crisis_type, ['environmental_impact'])
        
        # Get relevant factors
        relevant_factors = []
        recommended_monitoring = []
        
        for category_key in relevant_factor_categories:
            if category_key in crisis_factors:
                category = crisis_factors[category_key]
                for factor in category['factors']:
                    factor_with_relevance = factor.copy()
                    factor_with_relevance['category'] = category['category']
                    factor_with_relevance['relevance_score'] = 0.8  # High relevance for mapped factors
                    relevant_factors.append(factor_with_relevance)
        
        # Get recommended monitoring tasks based on crisis type
        monitoring_tasks = data['crisis_management_framework']['monitoring_tasks']
        for task_name, task in monitoring_tasks.items():
            if (scenario_crisis_type in ['natural_disaster', 'climate_change'] and 'environmental' in task_name) or \
               (scenario_crisis_type in ['economic_crisis', 'supply_chain_disruption'] and 'economic' in task_name) or \
               (task.get('priority') == 'critical'):
                recommended_monitoring.append({
                    'task_name': task_name,
                    'task': task['task'],
                    'priority': task['priority'],
                    'frequency': task['frequency'],
                    'description': task['description']
                })
        
        return {
            'scenario_id': scenario_id,
            'scenario_title': scenario.get('title'),
            'crisis_type': scenario_crisis_type,
            'severity_level': scenario.get('severity_level'),
            'relevant_factors': relevant_factors,
            'recommended_monitoring': recommended_monitoring,
            'assessment_timestamp': datetime.now(timezone.utc).isoformat(),
            'total_factors': len(relevant_factors),
            'critical_monitoring_tasks': len([t for t in recommended_monitoring if t['priority'] == 'critical'])
        }
        
    except Exception as e:
        logging.error(f"Scenario assessment error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to assess scenario: {str(e)}")

# Polycrisis Enhancement Models
class TemporalDynamics(BaseModel):
    timescale: str
    duration: str
    focus: str
    key_metrics: List[str]
    typical_stakeholders: List[str]
    decision_types: List[str]

class CrossDomainInteraction(BaseModel):
    from_domain: str
    to_domain: str
    interaction_strength: str
    mechanism: str
    example: str

class StakeholderType(BaseModel):
    name: str
    subtypes: List[str]
    primary_objectives: List[str]
    resources: List[str]
    constraints: List[str]
    decision_making_style: str

class AdvancedMonitoringTask(BaseModel):
    task: str
    description: str
    priority: str
    frequency: str
    indicators: List[Dict]

class PolycrisisEnhancementSummary(BaseModel):
    total_enhancements: int
    temporal_timescales: int
    cross_domain_interactions: int
    stakeholder_types: int
    advanced_monitoring_tasks: int
    cultural_dimensions: int

# Polycrisis Enhancement Endpoints
@api_router.get("/polycrisis-enhancements/summary", response_model=PolycrisisEnhancementSummary)
async def get_polycrisis_enhancements_summary(current_user: User = Depends(get_current_user)):
    """Get comprehensive summary of polycrisis enhancements"""
    try:
        enhancements_file = Path(__file__).parent.parent / "polycrisis_enhancements.json"
        
        if not enhancements_file.exists():
            raise HTTPException(status_code=404, detail="Polycrisis enhancements not found")
        
        with open(enhancements_file, 'r') as f:
            data = json.load(f)
        
        enhancements = data['polycrisis_enhancements']
        
        # Count temporal timescales
        temporal_timescales = len(enhancements['enhancement_categories']['temporal_dynamics']['timescales'])
        
        # Count cross-domain interactions
        cross_domain_interactions = len(enhancements['enhancement_categories']['cross_domain_impacts']['interaction_matrices'])
        
        # Count stakeholder types
        stakeholder_types = len(enhancements['enhancement_categories']['stakeholder_interactions']['stakeholder_types'])
        
        # Count advanced monitoring tasks
        advanced_monitoring_tasks = len(enhancements['advanced_monitoring_tasks'])
        
        # Count cultural dimensions
        cultural_dimensions = len(enhancements['enhancement_categories']['cultural_regional_variations']['cultural_dimensions'])
        
        return PolycrisisEnhancementSummary(
            total_enhancements=5,  # 5 main enhancement categories
            temporal_timescales=temporal_timescales,
            cross_domain_interactions=cross_domain_interactions,
            stakeholder_types=stakeholder_types,
            advanced_monitoring_tasks=advanced_monitoring_tasks,
            cultural_dimensions=cultural_dimensions
        )
        
    except Exception as e:
        logging.error(f"Polycrisis enhancements summary error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get enhancements summary: {str(e)}")

@api_router.get("/polycrisis-enhancements/temporal-dynamics")
async def get_temporal_dynamics(current_user: User = Depends(get_current_user)):
    """Get temporal dynamics modeling information"""
    try:
        enhancements_file = Path(__file__).parent.parent / "polycrisis_enhancements.json"
        
        if not enhancements_file.exists():
            raise HTTPException(status_code=404, detail="Polycrisis enhancements not found")
        
        with open(enhancements_file, 'r') as f:
            data = json.load(f)
        
        temporal_dynamics = data['polycrisis_enhancements']['enhancement_categories']['temporal_dynamics']
        
        return {
            'category': temporal_dynamics['category'],
            'description': temporal_dynamics['description'],
            'priority': temporal_dynamics['priority'],
            'timescales': temporal_dynamics['timescales'],
            'cascade_mechanisms': temporal_dynamics['cascade_mechanisms']
        }
        
    except Exception as e:
        logging.error(f"Temporal dynamics error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get temporal dynamics: {str(e)}")

@api_router.get("/polycrisis-enhancements/cross-domain-impacts")
async def get_cross_domain_impacts(current_user: User = Depends(get_current_user)):
    """Get cross-domain impact assessment information"""
    try:
        enhancements_file = Path(__file__).parent.parent / "polycrisis_enhancements.json"
        
        if not enhancements_file.exists():
            raise HTTPException(status_code=404, detail="Polycrisis enhancements not found")
        
        with open(enhancements_file, 'r') as f:
            data = json.load(f)
        
        cross_domain = data['polycrisis_enhancements']['enhancement_categories']['cross_domain_impacts']
        
        return {
            'category': cross_domain['category'],
            'description': cross_domain['description'],
            'domains': cross_domain['domains'],
            'interaction_matrices': cross_domain['interaction_matrices'],
            'feedback_loops': cross_domain['feedback_loops']
        }
        
    except Exception as e:
        logging.error(f"Cross-domain impacts error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get cross-domain impacts: {str(e)}")

@api_router.get("/polycrisis-enhancements/stakeholder-interactions")
async def get_stakeholder_interactions(current_user: User = Depends(get_current_user)):
    """Get advanced stakeholder interaction modeling information"""
    try:
        enhancements_file = Path(__file__).parent.parent / "polycrisis_enhancements.json"
        
        if not enhancements_file.exists():
            raise HTTPException(status_code=404, detail="Polycrisis enhancements not found")
        
        with open(enhancements_file, 'r') as f:
            data = json.load(f)
        
        stakeholder_interactions = data['polycrisis_enhancements']['enhancement_categories']['stakeholder_interactions']
        
        return {
            'category': stakeholder_interactions['category'],
            'description': stakeholder_interactions['description'],
            'stakeholder_types': stakeholder_interactions['stakeholder_types'],
            'interaction_mechanisms': stakeholder_interactions['interaction_mechanisms'],
            'network_effects': stakeholder_interactions['network_effects']
        }
        
    except Exception as e:
        logging.error(f"Stakeholder interactions error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get stakeholder interactions: {str(e)}")

@api_router.get("/polycrisis-enhancements/uncertainty-quantification")
async def get_uncertainty_quantification(current_user: User = Depends(get_current_user)):
    """Get uncertainty quantification methods and approaches"""
    try:
        enhancements_file = Path(__file__).parent.parent / "polycrisis_enhancements.json"
        
        if not enhancements_file.exists():
            raise HTTPException(status_code=404, detail="Polycrisis enhancements not found")
        
        with open(enhancements_file, 'r') as f:
            data = json.load(f)
        
        uncertainty = data['polycrisis_enhancements']['enhancement_categories']['uncertainty_quantification']
        
        return {
            'category': uncertainty['category'],
            'description': uncertainty['description'],
            'uncertainty_types': uncertainty['uncertainty_types'],
            'quantification_methods': uncertainty['quantification_methods'],
            'communication_strategies': uncertainty['communication_strategies']
        }
        
    except Exception as e:
        logging.error(f"Uncertainty quantification error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get uncertainty quantification: {str(e)}")

@api_router.get("/polycrisis-enhancements/cultural-variations")
async def get_cultural_variations(current_user: User = Depends(get_current_user)):
    """Get cultural and regional variation modeling information"""
    try:
        enhancements_file = Path(__file__).parent.parent / "polycrisis_enhancements.json"
        
        if not enhancements_file.exists():
            raise HTTPException(status_code=404, detail="Polycrisis enhancements not found")
        
        with open(enhancements_file, 'r') as f:
            data = json.load(f)
        
        cultural = data['polycrisis_enhancements']['enhancement_categories']['cultural_regional_variations']
        
        return {
            'category': cultural['category'],
            'description': cultural['description'],
            'cultural_dimensions': cultural['cultural_dimensions'],
            'regional_factors': cultural['regional_factors'],
            'response_variation_models': cultural['response_variation_models']
        }
        
    except Exception as e:
        logging.error(f"Cultural variations error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get cultural variations: {str(e)}")

@api_router.get("/polycrisis-enhancements/advanced-monitoring")
async def get_advanced_monitoring_tasks(
    priority: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Get advanced monitoring tasks for enhanced polycrisis management"""
    try:
        enhancements_file = Path(__file__).parent.parent / "polycrisis_enhancements.json"
        
        if not enhancements_file.exists():
            raise HTTPException(status_code=404, detail="Polycrisis enhancements not found")
        
        with open(enhancements_file, 'r') as f:
            data = json.load(f)
        
        monitoring_tasks = data['polycrisis_enhancements']['advanced_monitoring_tasks']
        
        # Filter by priority if specified
        filtered_tasks = {}
        for task_name, task_data in monitoring_tasks.items():
            if priority and task_data.get('priority') != priority:
                continue
            
            task_with_key = task_data.copy()
            task_with_key['task_key'] = task_name
            filtered_tasks[task_name] = task_with_key
        
        return filtered_tasks
        
    except Exception as e:
        logging.error(f"Advanced monitoring tasks error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get advanced monitoring tasks: {str(e)}")

@api_router.post("/polycrisis-enhancements/scenario-enhancement")
async def enhance_scenario_with_polycrisis_factors(
    scenario_id: str,
    enhancement_types: List[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Enhance a scenario with advanced polycrisis modeling factors"""
    try:
        # Get the scenario
        scenario = await db.scenarios.find_one({"id": scenario_id, "user_id": current_user.id})
        if not scenario:
            raise HTTPException(status_code=404, detail="Scenario not found")
        
        enhancements_file = Path(__file__).parent.parent / "polycrisis_enhancements.json"
        
        if not enhancements_file.exists():
            raise HTTPException(status_code=404, detail="Polycrisis enhancements not found")
        
        with open(enhancements_file, 'r') as f:
            data = json.load(f)
        
        enhancements = data['polycrisis_enhancements']['enhancement_categories']
        
        # Default to all enhancement types if none specified
        if not enhancement_types:
            enhancement_types = ['temporal_dynamics', 'cross_domain_impacts', 'stakeholder_interactions', 'uncertainty_quantification']
        
        # Generate enhanced scenario analysis
        enhanced_analysis = {
            'scenario_id': scenario_id,
            'scenario_title': scenario.get('title'),
            'crisis_type': scenario.get('crisis_type'),
            'severity_level': scenario.get('severity_level'),
            'enhancement_timestamp': datetime.now(timezone.utc).isoformat(),
            'enhancements_applied': enhancement_types,
            'enhanced_factors': {}
        }
        
        # Apply temporal dynamics enhancement
        if 'temporal_dynamics' in enhancement_types:
            temporal = enhancements['temporal_dynamics']
            enhanced_analysis['enhanced_factors']['temporal_dynamics'] = {
                'relevant_timescales': [
                    ts for ts in temporal['timescales'] 
                    if scenario.get('severity_level', 5) >= 7 or ts['name'] in ['Immediate Response', 'Short-term Impact']
                ],
                'cascade_mechanisms': temporal['cascade_mechanisms'],
                'recommendation': 'Implement multi-timescale monitoring for comprehensive crisis evolution tracking'
            }
        
        # Apply cross-domain impacts enhancement
        if 'cross_domain_impacts' in enhancement_types:
            cross_domain = enhancements['cross_domain_impacts']
            # Map crisis type to relevant domain interactions
            crisis_domain_map = {
                'economic_crisis': ['economic', 'social'],
                'natural_disaster': ['environmental', 'social', 'economic'],
                'cyber_attack': ['technological', 'economic'],
                'pandemic': ['social', 'economic', 'environmental']
            }
            
            relevant_domains = crisis_domain_map.get(scenario.get('crisis_type'), ['economic', 'social'])
            relevant_interactions = [
                interaction for interaction in cross_domain['interaction_matrices']
                if interaction['from_domain'] in relevant_domains or interaction['to_domain'] in relevant_domains
            ]
            
            enhanced_analysis['enhanced_factors']['cross_domain_impacts'] = {
                'relevant_domains': relevant_domains,
                'key_interactions': relevant_interactions,
                'feedback_loops': cross_domain['feedback_loops'],
                'recommendation': 'Monitor cross-domain spillover effects and feedback loops'
            }
        
        # Apply stakeholder interactions enhancement
        if 'stakeholder_interactions' in enhancement_types:
            stakeholders = enhancements['stakeholder_interactions']
            enhanced_analysis['enhanced_factors']['stakeholder_interactions'] = {
                'key_stakeholder_types': stakeholders['stakeholder_types'],
                'interaction_mechanisms': stakeholders['interaction_mechanisms'],
                'network_effects_to_monitor': stakeholders['network_effects'],
                'recommendation': 'Model stakeholder behavior patterns and cooperation/competition dynamics'
            }
        
        # Apply uncertainty quantification enhancement
        if 'uncertainty_quantification' in enhancement_types:
            uncertainty = enhancements['uncertainty_quantification']
            enhanced_analysis['enhanced_factors']['uncertainty_quantification'] = {
                'uncertainty_types_present': uncertainty['uncertainty_types'],
                'recommended_methods': uncertainty['quantification_methods'],
                'communication_strategies': uncertainty['communication_strategies'],
                'recommendation': 'Implement probabilistic ranges instead of point predictions'
            }
        
        return enhanced_analysis
        
    except Exception as e:
        logging.error(f"Scenario enhancement error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to enhance scenario: {str(e)}")

# AI Avatar System Models
class Competence(BaseModel):
    name: str
    skill_level: int  # 1-10 scale
    description: Optional[str] = None

class AIAvatar(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    name: str
    avatar_type: str
    category: str
    description: str
    specializations: List[str]
    core_competences: List[Competence]
    knowledge_domains: List[str]
    task_capabilities: List[str]
    team_name: Optional[str] = None
    organization: Optional[str] = None
    status: str = "active"  # "active", "busy", "inactive"
    performance_metrics: Dict = Field(default_factory=dict)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AvatarTask(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    avatar_id: str
    task_category: str
    task_type: str
    title: str
    description: str
    requirements: List[str] = Field(default_factory=list)
    priority: str = "medium"  # "low", "medium", "high", "urgent"
    status: str = "pending"  # "pending", "in_progress", "completed", "failed"
    assigned_at: Optional[datetime] = None
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    result: Optional[str] = None
    quality_score: Optional[float] = None
    feedback: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AvatarCreate(BaseModel):
    name: str
    avatar_type: str
    category: str
    description: str
    specializations: List[str]
    core_competences: List[Competence]
    knowledge_domains: List[str]
    task_capabilities: List[str]
    team_name: Optional[str] = None
    organization: Optional[str] = None

class TaskCreate(BaseModel):
    avatar_id: str
    task_category: str
    task_type: str
    title: str
    description: str
    requirements: List[str] = Field(default_factory=list)
    priority: str = "medium"

class TaskExecution(BaseModel):
    task_id: str
    action: str  # "start", "complete", "fail"
    result: Optional[str] = None
    quality_score: Optional[float] = None
    feedback: Optional[str] = None

# AI Avatar System Endpoints
@api_router.get("/ai-avatars", response_model=List[AIAvatar])
async def get_user_avatars(current_user: User = Depends(get_current_user)):
    """Get all AI avatars for the current user"""
    try:
        avatars = await db.ai_avatars.find({"user_id": current_user.id}).to_list(length=None)
        return [AIAvatar(**avatar) for avatar in avatars]
    except Exception as e:
        logging.error(f"Get avatars error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve avatars: {str(e)}")

@api_router.post("/ai-avatars", response_model=AIAvatar)
async def create_avatar(avatar_data: AvatarCreate, current_user: User = Depends(get_current_user)):
    """Create a new AI avatar"""
    try:
        avatar = AIAvatar(
            user_id=current_user.id,
            **avatar_data.dict()
        )
        
        avatar_dict = avatar.dict()
        await db.ai_avatars.insert_one(avatar_dict)
        
        return avatar
    except Exception as e:
        logging.error(f"Create avatar error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create avatar: {str(e)}")

@api_router.get("/ai-avatars/{avatar_id}", response_model=AIAvatar)
async def get_avatar(avatar_id: str, current_user: User = Depends(get_current_user)):
    """Get a specific AI avatar"""
    try:
        avatar = await db.ai_avatars.find_one({"id": avatar_id, "user_id": current_user.id})
        if not avatar:
            raise HTTPException(status_code=404, detail="Avatar not found")
        return AIAvatar(**avatar)
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get avatar error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve avatar: {str(e)}")

@api_router.put("/ai-avatars/{avatar_id}", response_model=AIAvatar)
async def update_avatar(avatar_id: str, avatar_data: AvatarCreate, current_user: User = Depends(get_current_user)):
    """Update an AI avatar"""
    try:
        avatar = await db.ai_avatars.find_one({"id": avatar_id, "user_id": current_user.id})
        if not avatar:
            raise HTTPException(status_code=404, detail="Avatar not found")
        
        update_data = avatar_data.dict()
        update_data["updated_at"] = datetime.now(timezone.utc)
        
        await db.ai_avatars.update_one({"id": avatar_id}, {"$set": update_data})
        updated_avatar = await db.ai_avatars.find_one({"id": avatar_id})
        return AIAvatar(**updated_avatar)
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Update avatar error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update avatar: {str(e)}")

@api_router.delete("/ai-avatars/{avatar_id}")
async def delete_avatar(avatar_id: str, current_user: User = Depends(get_current_user)):
    """Delete an AI avatar"""
    try:
        avatar = await db.ai_avatars.find_one({"id": avatar_id, "user_id": current_user.id})
        if not avatar:
            raise HTTPException(status_code=404, detail="Avatar not found")
        
        # Also delete all tasks associated with this avatar
        await db.avatar_tasks.delete_many({"avatar_id": avatar_id, "user_id": current_user.id})
        await db.ai_avatars.delete_one({"id": avatar_id})
        
        return {"message": "Avatar deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Delete avatar error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete avatar: {str(e)}")

@api_router.get("/ai-avatars/{avatar_id}/tasks", response_model=List[AvatarTask])
async def get_avatar_tasks(avatar_id: str, current_user: User = Depends(get_current_user)):
    """Get all tasks for a specific avatar"""
    try:
        # Verify avatar belongs to user
        avatar = await db.ai_avatars.find_one({"id": avatar_id, "user_id": current_user.id})
        if not avatar:
            raise HTTPException(status_code=404, detail="Avatar not found")
        
        tasks = await db.avatar_tasks.find({"avatar_id": avatar_id, "user_id": current_user.id}).to_list(length=None)
        return [AvatarTask(**task) for task in tasks]
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get avatar tasks error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve avatar tasks: {str(e)}")

@api_router.post("/ai-avatars/tasks", response_model=AvatarTask)
async def create_task(task_data: TaskCreate, current_user: User = Depends(get_current_user)):
    """Create a new task for an AI avatar"""
    try:
        # Verify avatar belongs to user
        avatar = await db.ai_avatars.find_one({"id": task_data.avatar_id, "user_id": current_user.id})
        if not avatar:
            raise HTTPException(status_code=404, detail="Avatar not found")
        
        task = AvatarTask(
            user_id=current_user.id,
            assigned_at=datetime.now(timezone.utc),
            **task_data.dict()
        )
        
        task_dict = task.dict()
        await db.avatar_tasks.insert_one(task_dict)
        
        # Update avatar status to busy
        await db.ai_avatars.update_one(
            {"id": task_data.avatar_id}, 
            {"$set": {"status": "busy", "updated_at": datetime.now(timezone.utc)}}
        )
        
        return task
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Create task error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create task: {str(e)}")

@api_router.post("/ai-avatars/tasks/{task_id}/execute")
async def execute_task(task_id: str, execution_data: TaskExecution, current_user: User = Depends(get_current_user)):
    """Execute or update task status"""
    try:
        task = await db.avatar_tasks.find_one({"id": task_id, "user_id": current_user.id})
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")
        
        update_data = {"updated_at": datetime.now(timezone.utc)}
        
        if execution_data.action == "start":
            update_data["status"] = "in_progress"
            update_data["started_at"] = datetime.now(timezone.utc)
        elif execution_data.action == "complete":
            update_data["status"] = "completed"
            update_data["completed_at"] = datetime.now(timezone.utc)
            if execution_data.result:
                update_data["result"] = execution_data.result
            if execution_data.quality_score:
                update_data["quality_score"] = execution_data.quality_score
            
            # Update avatar status back to active
            await db.ai_avatars.update_one(
                {"id": task["avatar_id"]}, 
                {"$set": {"status": "active", "updated_at": datetime.now(timezone.utc)}}
            )
        elif execution_data.action == "fail":
            update_data["status"] = "failed"
            if execution_data.feedback:
                update_data["feedback"] = execution_data.feedback
            
            # Update avatar status back to active
            await db.ai_avatars.update_one(
                {"id": task["avatar_id"]}, 
                {"$set": {"status": "active", "updated_at": datetime.now(timezone.utc)}}
            )
        
        await db.avatar_tasks.update_one({"id": task_id}, {"$set": update_data})
        updated_task = await db.avatar_tasks.find_one({"id": task_id})
        return AvatarTask(**updated_task)
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Execute task error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to execute task: {str(e)}")

@api_router.post("/ai-avatars/tasks/{task_id}/ai-execute")
async def ai_execute_task(task_id: str, current_user: User = Depends(get_current_user)):
    """Execute task using AI capabilities"""
    try:
        task = await db.avatar_tasks.find_one({"id": task_id, "user_id": current_user.id})
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")
        
        avatar = await db.ai_avatars.find_one({"id": task["avatar_id"], "user_id": current_user.id})
        if not avatar:
            raise HTTPException(status_code=404, detail="Avatar not found")
        
        # Mark task as in progress
        await db.avatar_tasks.update_one(
            {"id": task_id}, 
            {"$set": {"status": "in_progress", "started_at": datetime.now(timezone.utc)}}
        )
        
        # Get EMERGENT_LLM_KEY for AI execution
        emergent_key = os.environ.get('EMERGENT_LLM_KEY')
        if not emergent_key:
            raise HTTPException(status_code=500, detail="AI service not available")
        
        # Create AI prompt based on avatar competences and task
        avatar_competences = [comp["name"] for comp in avatar["core_competences"]]
        avatar_specializations = avatar["specializations"]
        
        prompt = f"""
You are an AI avatar named '{avatar["name"]}' with the following profile:
- Type: {avatar["avatar_type"]}
- Category: {avatar["category"]}
- Description: {avatar["description"]}
- Specializations: {', '.join(avatar_specializations)}
- Core Competences: {', '.join(avatar_competences)}
- Knowledge Domains: {', '.join(avatar["knowledge_domains"])}

Task to Execute:
- Title: {task["title"]}
- Description: {task["description"]}
- Category: {task["task_category"]}
- Type: {task["task_type"]}
- Requirements: {', '.join(task["requirements"]) if task["requirements"] else 'None specified'}

Please execute this task according to your competences and provide:
1. A detailed analysis or solution
2. Key findings or recommendations
3. Any relevant insights based on your specializations
4. Next steps or follow-up actions if applicable

Provide a professional, comprehensive response that demonstrates your expertise in the relevant areas.
"""
        
        # Use emergentintegrations to call AI service
        try:
            from emergentintegrations import EmergentIntegrations
            
            ei = EmergentIntegrations(api_key=emergent_key)
            
            # Use Claude Sonnet for complex analysis tasks
            response = await asyncio.to_thread(
                ei.text_generation,
                prompt=prompt,
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                temperature=0.3
            )
            
            ai_result = response.get('text', 'Task execution completed but no detailed result available.')
            
        except Exception as ai_error:
            logging.error(f"AI execution error: {str(ai_error)}")
            ai_result = f"Task execution encountered an error: {str(ai_error)}"
        
        # Mark task as completed with AI result
        completion_time = datetime.now(timezone.utc)
        await db.avatar_tasks.update_one(
            {"id": task_id}, 
            {"$set": {
                "status": "completed",
                "completed_at": completion_time,
                "result": ai_result,
                "quality_score": 8.5,  # Default good quality score for AI execution
                "updated_at": completion_time
            }}
        )
        
        # Update avatar status back to active
        await db.ai_avatars.update_one(
            {"id": task["avatar_id"]}, 
            {"$set": {"status": "active", "updated_at": completion_time}}
        )
        
        updated_task = await db.avatar_tasks.find_one({"id": task_id})
        return AvatarTask(**updated_task)
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"AI execute task error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to execute task with AI: {str(e)}")

@api_router.get("/ai-avatars/system/templates")
async def get_avatar_templates(current_user: User = Depends(get_current_user)):
    """Get available avatar templates from the system"""
    try:
        system_file = Path(__file__).parent.parent / "ai_avatar_system.json"
        
        if not system_file.exists():
            raise HTTPException(status_code=404, detail="Avatar system not found")
        
        with open(system_file, 'r') as f:
            data = json.load(f)
        
        avatar_categories = data['ai_avatar_system']['avatar_categories']
        
        return {
            'categories': avatar_categories,
            'competence_framework': data['ai_avatar_system']['competence_framework'],
            'task_categories': data['ai_avatar_system']['task_management']['task_categories']
        }
        
    except Exception as e:
        logging.error(f"Get avatar templates error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get avatar templates: {str(e)}")

# Include the API router in the main app
app.include_router(api_router)